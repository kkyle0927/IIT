"""
PhAI Studio 통합 기술사양서 Word(.docx) 생성기.

Physical AI 플랫폼 — π0 Dual-System + Human Expert 프레임워크 기반
통합 기술 사양 문서.

Usage:
    python docs/xm-studio/gen_phai_studio_spec.py

Dependencies:
    pip install python-docx pyyaml
"""
import sys
import os

BASE_DIR = r"C:\AGR_HyundoKim\02_PhAI\PhAI_Web\Docs\Plan\PhAI_Studio_기획"
sys.path.insert(0, BASE_DIR)
sys.stdout.reconfigure(encoding="utf-8")

from gen_docx_v2_base import (
    setup_document, add_title_page, add_info_cards, add_table, add_code_block,
    add_body, add_bullet, add_ref, FONT_KR,
    BG_INFO_BLUE, BG_INFO_GRAY, BG_INFO_GREEN, BG_INFO_AMBER,
)

F = FONT_KR
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def _add_toc(doc, items, lang="kr"):
    from docx.shared import Pt
    from gen_docx_v2_base import CLR_GRAY
    doc.add_heading("목 차" if lang == "kr" else "Table of Contents", level=1)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("   " + item)
        r.font.name = F
        r.font.size = Pt(10)
        r.font.color.rgb = CLR_GRAY
    doc.add_page_break()


# ============================================================
# Korean Version
# ============================================================
def build_phai_studio_spec_kr(doc):
    """PhAI Studio 통합 기술사양서 (한국어)."""

    # === TITLE PAGE ===
    add_title_page(doc,
        "PhAI Studio\n통합 기술사양서",
        "Physical AI 플랫폼 — π0 Dual-System + Human Expert 프레임워크 기반",
        "v1.0  ·  2026-03-12  ·  Angel Robotics"
    )

    add_info_cards(doc, [
        ("적용 대상", "Physical AI 통합 플랫폼", BG_INFO_BLUE),
        ("핵심 철학", "π0 System 1/2 + Human Expert", BG_INFO_GRAY),
        ("기능 수", "6개 모듈 · 45개 기능", BG_INFO_GREEN),
        ("아키텍처", "웹 프론트엔드 + 로컬 에이전트", BG_INFO_GREEN),
    ])
    add_info_cards(doc, [
        ("시뮬레이션", "MuJoCo WASM + MJX", BG_INFO_BLUE),
        ("AI 학습", "RL/LfD + SLURM GPU", BG_INFO_AMBER),
        ("데이터", "1kHz 무손실 + HDF5", BG_INFO_GRAY),
        ("상태", "기획 확정, 개발 준비 중", BG_INFO_GRAY),
    ])

    toc_items = [
        "0. 기획자를 위한 핵심 기술 용어 사전",
        "1. 비전 & 철학",
        "2. 기능 체계 (6개 모듈, 45개 기능)",
        "3. Best Practice 기술 선정",
        "4. 시스템 아키텍처",
        "5. 아키텍처 결정 사항",
        "6. Phase 로드맵",
    ]
    _add_toc(doc, toc_items, "kr")

    # ================================================================
    # Part 0. 기술 용어 사전
    # ================================================================
    doc.add_heading("0. 기획자를 위한 핵심 기술 용어 사전", level=1)
    add_body(doc,
        "PhAI Studio 기획에 등장하는 기술 용어를 비유와 함께 설명합니다. "
        "전문 지식이 없는 기획자도 이해할 수 있도록 일상적인 비유를 첨부하였습니다."
    )
    add_table(doc,
        ["용어", "비유", "설명"],
        [
            ["WebAssembly (WASM)", "브라우저 안의 게임 엔진",
             "C/C++ 코드를 브라우저에서 실행할 수 있게 변환하는 기술. 네이티브 대비 ~55% 속도. "
             "MuJoCo 물리엔진을 브라우저에서 돌리는 핵심 기술"],
            ["WebSocket", "전화 통화",
             "서버와 브라우저 간 실시간 양방향 통신. HTTP가 '편지'라면 WebSocket은 '전화' — "
             "연결 유지, 데이터 즉시 전달"],
            ["Web Serial API", "브라우저 USB 포트",
             "Chrome 브라우저에서 USB 장치에 직접 접근. 별도 설치 없이 USB CDC 통신 가능"],
            ["로컬 에이전트", "프린터 드라이버",
             "PC에 설치하는 백그라운드 서비스(UI 없음). USB 통신, HDF5 저장, FW 업로드 프록시 역할. "
             "데스크톱 앱(GUI)이 아님. 사용자가 보고 조작하는 모든 화면은 웹 브라우저"],
            ["HDF5", "과학자의 엑셀",
             "대용량 수치 데이터를 효율적으로 저장하는 파일 형식. AI 학습 표준. "
             "1kHz×100채널 데이터 고속 읽기/쓰기"],
            ["URDF", "로봇 설계도 XML",
             "로봇의 링크(뼈), 조인트(관절), 질량, 관성을 기술하는 XML 파일. "
             "시뮬레이션과 시각화의 입력"],
            ["MuJoCo", "로봇 전용 물리엔진",
             "DeepMind 관리 로봇 역학 시뮬레이터. 접촉, 마찰, 관절 동역학 정밀 계산. "
             "Unity가 '게임용'이면 MuJoCo는 '연구/양산용'"],
            ["MJX", "MuJoCo GPU 터보",
             "MuJoCo를 GPU에서 수천 개 병렬 실행하는 확장. AI 대규모 학습 필수"],
            ["ONNX", "AI 모델 번역기",
             "PyTorch로 학습한 AI 모델을 다른 환경(브라우저, STM32)에서 실행 가능하게 변환하는 표준 포맷"],
            ["STM32Cube.AI", "AI를 칩에 굽는 도구",
             "ONNX 모델을 STM32 마이크로컨트롤러에서 실행 가능한 C 코드로 자동 변환하는 ST 공식 도구"],
            ["Sim-to-Real", "시뮬 → 현실 이식",
             "시뮬레이션에서 학습/검증한 제어기를 실제 로봇에 올려 동작시키는 과정"],
            ["SLURM", "GPU 예약 시스템",
             "여러 사용자가 GPU 서버를 공유할 때 순번 관리와 자원 배분을 해주는 스케줄러"],
            ["Monaco Editor", "VS Code의 심장",
             "VS Code의 코드 편집 엔진 브라우저 버전. 자동완성, 구문 강조, 에러 표시 등 IDE 기능 웹 제공"],
            ["Three.js", "브라우저 3D 엔진",
             "웹 브라우저에서 3D 그래픽 렌더링. URDF 로봇 모델 시각화에 사용"],
            ["Docker", "가상 컴퓨터 상자",
             "개발 환경을 통째로 패키징하여 어디서든 동일하게 실행. "
             "클라우드 빌드 서버에서 GCC 컴파일러 운영에 사용"],
            ["π0 (Pi-Zero)", "로봇 범용 두뇌",
             "Physical Intelligence(2024)의 로봇 파운데이션 모델. "
             "3B 파라미터 VLM 기반으로 68가지 조작 작업을 하나의 모델로 수행"],
            ["VLA", "보고, 듣고, 움직이는 AI",
             "Vision-Language-Action. 카메라(Vision) + 자연어(Language) → "
             "로봇 동작(Action) 생성 통합 모델"],
            ["System 1 / System 2", "반사 vs 숙고",
             "카너먼 인지 심리학 이론의 로봇 적용. "
             "System 1 = 빠른 반사 제어(1kHz MCU), System 2 = 느린 고차원 추론(AI 서버)"],
            ["Shared Autonomy", "운전 교습",
             "인간과 로봇이 제어권을 공유. "
             "초보 운전자(로봇)와 교관(전문가)처럼 점진적으로 자율성 확대"],
            ["Human-in-the-Loop", "사람이 루프 안에",
             "AI 학습/실행 과정에 인간이 지속 참여. "
             "수집 → 학습 → 평가 → 교정 모든 단계에 전문가 개입"],
        ], first_col_bold=True
    )

    # ================================================================
    # Part 1. 비전 & 철학
    # ================================================================
    doc.add_heading("1. 비전 & 철학", level=1)

    # 1.1
    doc.add_heading("1.1  π0에서 배우는 Dual-System 로봇 지능", level=2)
    add_body(doc,
        "Physical Intelligence가 2024년 발표한 π0(Pi-Zero)는 로봇을 위한 최초의 "
        "범용 정책(Generalist Policy) 모델이다. 3B 파라미터 Vision-Language-Action(VLA) "
        "아키텍처로, 빨래 개기, 테이블 정리 등 68가지 작업을 하나의 모델로 수행한다."
    )
    add_body(doc,
        "π0의 핵심 통찰은 카너먼(Daniel Kahneman)의 "
        "'생각에 관한 생각(Thinking, Fast and Slow)'에서 차용한 "
        "Dual-System 프레임워크이다:"
    )
    add_table(doc,
        ["구분", "System 1 (Fast)", "System 2 (Slow)"],
        [
            ["카너먼 비유", "직관적, 빠른 판단", "분석적, 느린 추론"],
            ["로봇 적용", "반사적 모터 제어 (1kHz)", "고차원 작업 계획, 환경 이해"],
            ["실행 위치", "MCU (STM32, BareMetal)", "AI 서버 / Jetson (VLA 모델)"],
            ["예시", "관절 토크 PD 제어, 마찰 보상", "'이 물건을 저기 놓아라' 이해"],
            ["주기", "1ms (1kHz)", "20~100ms (10~50Hz)"],
        ], first_col_bold=True
    )
    add_ref(doc,
        "[Ref] π0: A Vision-Language-Action Flow Model for General Robot Control "
        "— Physical Intelligence, 2024.10"
    )
    add_ref(doc,
        "[Ref] Thinking Fast and Slow in Human and Machine Intelligence — ACM CACM, 2025"
    )

    # 1.2
    doc.add_heading("1.2  ONE PhAI의 차별점: 인간 전문가를 시스템에 편입", level=2)
    add_body(doc,
        "π0는 System 1 + System 2의 '로봇 자체 지능'에 집중한다. "
        "ONE PhAI는 여기에 인간 전문가(Human Expert)를 시스템의 핵심 구성원으로 편입시킨다."
    )
    add_body(doc, "핵심 철학: '대체(Replace)가 아닌 보조(Augment)'", bold=True)
    add_body(doc,
        "로봇은 전문가가 가르친 범위 내에서 자율적으로 동작한다. "
        "예외 상황, 새로운 환경, 안전 판단은 여전히 인간이 담당한다. "
        "이는 제조업 현장에서 숙련공의 노하우를 보존하면서 생산성을 높이고, "
        "의료 현장에서 전문 치료사의 판단을 존중하면서 24시간 보조를 가능하게 하는 접근이다."
    )
    add_body(doc,
        "또한, 인간이 직접 작업하기 어려운 환경(고온, 분진, 방사선, 밀폐 공간 등)에서도 "
        "전문가의 시범을 학습한 로봇이 전문가 수준의 동작을 자율적으로 수행할 수 있다."
    )
    add_table(doc,
        ["역할", "π0 프레임워크", "ONE PhAI 프레임워크"],
        [
            ["System 1 (Fast)", "반사적 모터 제어", "동일 — MCU 1kHz 제어"],
            ["System 2 (Slow)", "VLA 고차원 추론", "동일 — AI 서버 학습/추론"],
            ["Human Expert", "학습 데이터 제공자 (수동적)",
             "시스템의 능동적 구성원\n(교시 → 교정 → 평가 → 감독)"],
            ["진화 방향", "데이터 확장 → 범용화", "전문가 협업 → 점진적 자율성 확대"],
        ], first_col_bold=True
    )

    # 1.3
    doc.add_heading("1.3  3단계 자율성 진화 모델", level=2)
    add_body(doc,
        "인간과 로봇의 역할은 고정이 아니라 시간에 따라 진화한다. "
        "PhAI Studio는 이 3단계 진화를 지원하는 도구 체계를 제공한다."
    )
    add_table(doc,
        ["단계", "인간 역할", "로봇 역할", "Studio 역할"],
        [
            ["Stage 1\nHuman-Dominant\n(인간 주도)",
             "직접 조작 100%\n시범 시연",
             "따라하기 (모방)\n수동적 학습",
             "궤적 녹화 (E7)\nDigital Twin (E8)\n데이터 수집 (D1)"],
            ["Stage 2\nShared Autonomy\n(공유 자율)",
             "감독 + 교정\n품질 평가",
             "기본 자율 동작\n+ 인간 보정 시 학습",
             "LfD/DAgger 학습 (T4)\n정책 평가 (E10)\n점진적 개선"],
            ["Stage 3\nRobot-Autonomous\n(로봇 자율)",
             "모니터링\n예외 처리",
             "전문가 수준 자율\n자체 판단",
             "성능 대시보드\n안전 경계 감시 (E11)\n이상 감지 알림"],
        ], first_col_bold=True
    )

    # 1.4
    doc.add_heading("1.4  적용 시나리오", level=2)
    add_table(doc,
        ["현장", "인간이 어려운 환경", "전문가 데이터", "Studio 역할"],
        [
            ["의료 재활 (H10)",
             "환자마다 다른 보행 패턴\n24시간 보조 필요",
             "물리치료사의\n보조 토크 패턴",
             "치료사 시범 녹화\n→ 개인화 학습"],
            ["제조 조립",
             "고온/분진/소음 환경\n장시간 반복 작업",
             "숙련공의\n손동작 + 힘 제어",
             "숙련공 시범\n→ 조립 자동화"],
            ["위험 검사",
             "고소/밀폐/방사선\n인명 위험 환경",
             "검사 전문가의\n판단 기준",
             "전문가 조작\n→ 자율 검사"],
            ["보행 보조",
             "장시간 보행\n계단/경사 적응",
             "보행 전문가의\n보조 전략",
             "전문가 데이터\n→ 적응형 보조"],
        ], first_col_bold=True
    )

    # 1.5
    doc.add_heading("1.5  PhAI Studio의 역할 매핑", level=2)
    add_body(doc,
        "PhAI Studio의 6개 모듈은 각각 이 철학의 특정 측면을 실현한다:"
    )
    add_table(doc,
        ["모듈", "철학적 역할", "핵심 질문 (Why)"],
        [
            ["Module 1\nActuator",
             "System 1의 물리적 기반\n구동기 투명성 확보",
             "구동기 투명성이 없으면 전문가의 의도가\n로봇에 왜곡되어 전달된다"],
            ["Module 2\nSensor",
             "전문가 데이터 수집 통로\n1kHz 무손실",
             "1kHz 무손실이 아니면 전문가 동작의\n미세한 뉘앙스가 소실된다"],
            ["Module 3\nData",
             "데이터 시각화 + 관리\nDigital Twin",
             "전문가가 자신의 동작을 확인할 수 없으면\n효과적인 교시가 불가능하다"],
            ["Module 4\nAI + Sim",
             "System 2 학습\nRL/LfD/시뮬레이션",
             "시뮬레이터 없이 학습하면 위험 환경에서\n실물로만 시도해야 한다"],
            ["Module 5\nHuman",
             "Shared Autonomy 인터페이스\n교시/교정/평가",
             "전문가가 개입할 수 없으면\n로봇의 자율성이 점진적으로 향상될 수 없다"],
            ["Module 6\nPlatform",
             "다수 사용자 인프라\n확장성 + 적응 품질",
             "한 명만 쓸 수 있는 도구는\n플랫폼이 될 수 없다"],
        ], first_col_bold=True
    )

    # ================================================================
    # Part 2. 기능 체계
    # ================================================================
    doc.add_heading("2. 기능 체계 (6개 모듈, 45개 기능)", level=1)
    add_body(doc,
        "ONE PhAI의 5대 모듈(구동기, 센서, 데이터, AI, 인간 상호작용)에 "
        "플랫폼 인프라를 더한 6개 모듈에 PhAI Studio의 45개 기능을 매핑한다."
    )

    # 2.1
    doc.add_heading("2.1  Module 1: Actuator Commissioning (구동기 셋업)", level=2)
    add_table(doc,
        ["ID", "기능명", "설명", "현재 상태"],
        [
            ["A1", "Actuator System Identification",
             "구동기의 전기적/기계적 파라미터를 측정하여 수학적 모델 추출. "
             "토크 입출력 1:1 투명성 확보의 핵심",
             "SAM_GUI (MATLAB)"],
            ["A2", "Friction Model Calibration",
             "쿨롱/점성/스트리벡 마찰 모델 피팅 및 마찰 보상기 설계",
             "SAM_GUI (MATLAB)"],
            ["A3", "Feedforward Controller Design",
             "식별된 모델 기반 피드포워드 제어기 자동 설계",
             "SAM_GUI (MATLAB)"],
            ["A4", "Frequency Response Analysis",
             "보드 선도(Bode Plot) 기반 대역폭 분석 및 검증",
             "SAM_GUI (MATLAB)"],
            ["A5", "Actuator Parameter Database",
             "식별 완료된 구동기 파라미터를 DB화. "
             "MD Flash에 저장되어 Sim-to-Real 시 자동 적용",
             "미구현"],
        ], first_col_bold=True
    )

    # 2.2
    doc.add_heading("2.2  Module 2: Sensor Integration & Data Acquisition", level=2)
    add_table(doc,
        ["ID", "기능명", "설명"],
        [
            ["S1", "Real-time Data Streaming (1kHz)",
             "USB CDC / CAN-FD 기반 1kHz 데이터 스트리밍. "
             "유실 감지 + 시퀀스 번호 검증"],
            ["S2", "Data Integrity Monitoring",
             "패킷 유실률 실시간 표시, 유실 시 알림 + 보간(Cubic Spline) 후처리"],
            ["S3", "Multi-Device Management",
             "복수 모듈(XM, MD, SM) 동시 연결 및 데이터 동기화"],
            ["S4", "Device Connection Manager",
             "디바이스 자동 탐지, 연결 상태 모니터링, 재연결"],
        ], first_col_bold=True
    )

    # 2.3
    doc.add_heading("2.3  Module 3: Data Management & Visualization", level=2)
    add_table(doc,
        ["ID", "기능명", "설명"],
        [
            ["D1", "Lossless Data Recording",
             "1kHz 전 채널 데이터를 HDF5로 무손실 저장 (타임스탬프, 메타데이터 포함)"],
            ["D2", "Real-time Signal Visualization",
             "1kHz 데이터의 실시간 그래프 (최소 60FPS 렌더링, LTTB 다운샘플링)"],
            ["D3", "3D Robot Visualization",
             "URDF 로봇 모델의 실시간 관절 상태 시각화 (60FPS)"],
            ["D4", "Data Post-processing Pipeline",
             "필터링(LPF, Butterworth), 미분/적분, FFT, 통계 분석"],
            ["D5", "Dataset Version Control",
             "학습용 데이터셋의 버전 관리, 메타데이터 태깅, 검색"],
            ["D6", "Data Export & Format Conversion",
             "HDF5, CSV, ROS Bag, MCAP 등 다양한 형식 변환"],
        ], first_col_bold=True
    )

    # 2.4
    doc.add_heading("2.4  Module 4: Dynamics Simulation & AI Training", level=2)
    add_table(doc,
        ["ID", "기능명", "설명"],
        [
            ["T1", "Interactive Dynamics Simulator",
             "URDF 로봇의 동역학 시뮬레이션. "
             "마우스 드래그로 힘/토크 적용, 실시간 물리 반응"],
            ["T2", "Control Algorithm Sandbox",
             "사용자 작성 제어 코드(C/Python)를 시뮬레이터에서 즉시 실행"],
            ["T3", "RL Training Pipeline",
             "GPU 서버에서 MuJoCo MJX 기반 병렬 강화학습 (PPO/SAC)"],
            ["T4", "Imitation Learning / LfD Pipeline",
             "시뮬레이터/실물 조작 데이터로부터 모방학습 (ACT, Diffusion Policy)"],
            ["T5", "Sim-to-Real Deployment",
             "학습된 정책을 ONNX → STM32Cube.AI → XM10/MD FW로 배포"],
            ["T6", "Model Evaluation & Benchmarking",
             "학습 결과 정량 평가 (MSE, MAE, R², Success Rate)"],
            ["T7", "Domain Randomization Manager",
             "시뮬레이션 파라미터(마찰, 질량, 지연) 무작위화 관리"],
        ], first_col_bold=True
    )

    # 2.5
    doc.add_heading("2.5  Module 5: Development & Human Interaction", level=2)
    add_table(doc,
        ["ID", "기능명", "설명"],
        [
            ["E1", "Web IDE — Sandbox Mode",
             "예제 갤러리에서 코드 선택, 수정, 즉시 빌드/실행"],
            ["E2", "Web IDE — Direct Coding Mode",
             "빈 프로젝트에서 XM_API 기반 C 코드 자유 작성"],
            ["E3", "Web IDE — LLM-Assisted Coding",
             "AI(Claude/GPT) 기반 코드 생성, XM_API 제약 내 (Phase 2)"],
            ["E4", "Cloud Compilation Service",
             "Docker + arm-none-eabi-gcc로 서버에서 크로스 컴파일"],
            ["E5", "Firmware Upload (OTA/USB)",
             "packaged.bin을 Web Serial 또는 에이전트 경유 XM10 업로드"],
            ["E6", "Python Scripting Environment",
             "데이터 분석, 시뮬레이션 제어, 학습 스크립트용 Python 환경"],
            ["E7", "Expert Demonstration Recorder",
             "전문가가 실물/시뮬에서 로봇을 조작하며 궤적 녹화 (1kHz)"],
            ["E8", "Live Digital Twin",
             "실물 로봇의 1kHz 데이터를 시뮬 URDF 모델에 실시간 미러링"],
            ["E9", "Trajectory Replay & LfD",
             "녹화된 궤적을 시뮬에서 재생, 편집, LfD 학습 데이터로 변환"],
            ["E10", "Policy Evaluation Dashboard",
             "학습된 정책의 동작을 전문가가 시각적으로 평가, 승인/거절/교정"],
            ["E11", "Safety Boundary Editor",
             "관절 한계, 최대 토크, 충돌 영역 등 안전 경계를 시각적으로 설정"],
        ], first_col_bold=True
    )

    # 2.6
    doc.add_heading("2.6  Module 6: Platform Infrastructure", level=2)
    add_table(doc,
        ["ID", "기능명", "설명"],
        [
            ["P1", "User Authentication & Authorization",
             "사용자 인증, 프로젝트 권한 관리"],
            ["P2", "GPU Resource Scheduler",
             "SLURM 기반 다수 사용자 GPU 자원 배분 + 작업 큐"],
            ["P3", "Adaptive Quality System",
             "클라이언트 PC 사양/네트워크에 따라 시각화 품질 자동 조절"],
            ["P4", "Session Isolation",
             "사용자별 독립 시뮬레이션/빌드 환경"],
            ["P5", "Telemetry & Usage Analytics",
             "플랫폼 사용량, 성능, 에러 모니터링"],
            ["P6", "Landing & Onboarding Page",
             "비로그인 사용자 소개 페이지. PhAI Studio 비전, 기능 갤러리, "
             "데모 영상, 로봇 3D 프리뷰"],
            ["P7", "Robot Model Library",
             "URDF 로봇 모델 갤러리. PhAI-X1(H10), 향후 4족/매니퓰레이터/휴머노이드. "
             "접속 시 자동 식별"],
        ], first_col_bold=True
    )

    # ================================================================
    # Part 3. Best Practice 기술 선정
    # ================================================================
    doc.add_heading("3. Best Practice 기술 선정", level=1)
    add_body(doc,
        "각 기능에 대한 기술 선택을 근거와 레퍼런스와 함께 정리한다. "
        "모든 결정은 양산 품질 수준의 실현 가능성을 최우선으로 검증하였다."
    )

    # 3.1
    doc.add_heading("3.1  [A1~A5] 구동기 셋업 — MATLAB → Python/Web 마이그레이션", level=2)
    add_body(doc,
        "AS-IS: SAM_GUI (MATLAB App) — do_sys_id_analysis.m, do_friction_model_fitting.m 등",
        bold=True
    )
    add_body(doc, "TO-BE: Python 백엔드 + 웹 프론트엔드", bold=True)
    add_table(doc,
        ["항목", "기술 선정", "근거"],
        [
            ["수치 연산", "SciPy + NumPy (서버)",
             "MATLAB lsqcurvefit → scipy.optimize.curve_fit, fft → scipy.fft 1:1 대응"],
            ["제어 공학", "python-control",
             "MATLAB Control System Toolbox 대체. 보드 선도, 전달함수, 상태공간"],
            ["실시간 그래프", "uPlot (프론트) + Matplotlib (리포트)",
             "실시간 → uPlot, 보고서 → Matplotlib PDF"],
            ["통신", "PySerial (에이전트)",
             "MATLAB serialport() 대체. USB CDC 직접 제어"],
            ["파라미터 DB", "FastAPI + SQLite/PostgreSQL",
             "식별 결과 저장, 로봇별 조회, Sim-to-Real 자동 적용"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] SysIdentPy — sysidentpy.org (비선형 시스템 식별 Python 라이브러리)")
    add_ref(doc, "[Ref] motulator — aalto-electric-drives.github.io (모터 드라이브 시뮬레이션)")

    # 3.2
    doc.add_heading("3.2  [S1~S4] 데이터 취득 — 1kHz 무손실", level=2)
    add_table(doc,
        ["항목", "기술 선정", "근거"],
        [
            ["USB 통신", "로컬 에이전트 (PySerial) Primary\n+ Web Serial Lite Fallback",
             "PySerial: OS 네이티브, 1kHz 안정\nWeb Serial: 에이전트 없는 과도기 대응"],
            ["프로토콜", "PhAI V2.2 (COBS+CRC16, SOF=0xAA)",
             "검증됨. 시퀀스 번호 추가로 유실 감지"],
            ["유실 감지", "시퀀스 번호 (uint16 wrapping)",
             "매 패킷 증가 번호. 번호 불연속 = 유실"],
            ["보간", "Cubic Spline (scipy.interpolate)",
             "1~2 샘플 유실 시 보간. 3+ 연속 유실 시 경고"],
            ["저장", "HDF5 (h5py)",
             "로봇 데이터 업계 표준 (Foxglove, ALOHA, HumanoidBench)"],
        ], first_col_bold=True
    )

    # 3.3
    doc.add_heading("3.3  [D1~D6] 시각화 — 1kHz 수집 중 60FPS 보장", level=2)
    add_body(doc, "핵심 설계 원칙: 수집 스레드와 시각화 스레드 완전 분리", bold=True)
    add_code_block(doc,
        "에이전트 (Python):\n"
        "  Thread 1: USB 수신 → HDF5 저장 (1kHz, 블로킹 없음)\n"
        "  Thread 2: WebSocket → 브라우저에 다운샘플링 데이터 전송 (60Hz)\n"
        "\n"
        "브라우저:\n"
        "  WebSocket 수신 → uPlot 그래프 (60FPS) + Three.js URDF (60FPS)\n"
        "\n"
        "핵심: 시각화가 느려져도 저장은 영향 없음 (독립 스레드)"
    )
    add_table(doc,
        ["항목", "기술 선정", "근거"],
        [
            ["실시간 그래프", "uPlot",
             "CPU 10%, RAM 12MB, 16만 포인트 25ms. 타 라이브러리 대비 압도적 성능"],
            ["3D 시각화", "Three.js + urdf-loader",
             "168KB. URDF 직접 로드. React Three Fiber로 선언적 3D"],
            ["다운샘플링", "LTTB (Largest Triangle Three Buckets)",
             "시각적 형태 보존하면서 포인트 수 감소. 1kHz → 60Hz"],
            ["후처리", "SciPy (서버)",
             "Butterworth LPF, FFT, 통계 — MATLAB 동등 기능"],
        ], first_col_bold=True
    )

    # 3.4
    doc.add_heading("3.4  [T1~T7] 시뮬레이션 & AI 학습", level=2)
    add_table(doc,
        ["항목", "기술 선정", "근거"],
        [
            ["물리 엔진 (인터랙티브)", "MuJoCo WASM (브라우저)",
             "네트워크 지연 0. mujoco_wasm, mjswan 검증"],
            ["물리 엔진 (대규모 학습)", "MuJoCo MJX (서버 GPU)",
             "단일 GPU 65만 step/s. MuJoCo Playground RSS 2025 데모"],
            ["마우스 인터랙션", "MuJoCo perturbation API",
             "Ctrl+드래그=힘, Ctrl+우클릭=토크. MuJoCo 네이티브"],
            ["3D 모델 형식", "URDF",
             "로봇 업계 표준. MuJoCo URDF 로더 + Three.js urdf-loader"],
            ["RL 프레임워크", "Stable-Baselines3 (PPO/SAC)",
             "PyTorch 기반 표준. MuJoCo Playground 호환"],
            ["모방 학습", "LeRobot (Hugging Face)",
             "ACT, Diffusion Policy 구현. HDF5 데이터셋 호환"],
            ["점진적 학습", "DAgger (Dataset Aggregation)",
             "전문가 교정 데이터를 기존 정책에 점진 통합"],
            ["분산 학습", "SLURM + ClusterGym 또는 Ray",
             "32+ 노드 선형 확장. 비동기 패턴"],
            ["Sim-to-Real", "ONNX → STM32Cube.AI",
             "STM32H743에서 소형 MLP: <1ms 추론"],
            ["브라우저 AI 추론", "ONNX Runtime Web",
             "학습된 정책을 브라우저에서 직접 실행"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] MuJoCo WASM — github.com/zalo/mujoco_wasm")
    add_ref(doc, "[Ref] mjswan — github.com/ttktjmt/mjswan (MuJoCo WASM + ONNX RT Web + Three.js)")
    add_ref(doc, "[Ref] MuJoCo Playground — playground.mujoco.org (RSS 2025 Best Demo)")
    add_ref(doc, "[Ref] LeRobot — github.com/huggingface/lerobot")
    add_ref(doc, "[Ref] ClusterGym — github.com/rodlaf/ClusterGym (MuJoCo + SLURM)")

    # 3.5
    doc.add_heading("3.5  [E1~E11] 개발 환경 & 인간 상호작용", level=2)
    add_table(doc,
        ["항목", "기술 선정", "근거"],
        [
            ["코드 에디터", "Monaco Editor",
             "VS Code 심장. 60+ 언어, IntelliSense. XM_API JSON Schema로 커스텀 자동완성"],
            ["클라우드 빌드", "Docker + arm-none-eabi-gcc\n+ libXM_Full.a",
             "사용자 코드만 컴파일 (3~8초 목표)"],
            ["FW 업로드 (Phase 0)", "Web Serial (ESP Web Tools 패턴)",
             "수백만 대 양산 검증"],
            ["FW 업로드 (Phase 1+)", "로컬 에이전트 (PySerial)",
             "안정성 강화, 백그라운드 동작"],
            ["LLM 코딩", "API 연동 (Claude/GPT)\n+ RAG (XM_API 문서)",
             "Phase 2 이후. XM_API 제약 내 코드 생성"],
            ["시범 녹화", "E7: WebSocket 궤적 스트림\n+ HDF5 저장",
             "1kHz 관절각 + 힘/토크 + 타임스탬프 기록"],
            ["Digital Twin", "E8: 에이전트 60Hz 데이터\n→ Three.js URDF 애니메이션",
             "실물 동작을 시뮬에서 미러링"],
            ["궤적 재생", "E9: HDF5 → uPlot 타임라인\n+ Three.js 3D 재생",
             "편집/잘라내기/이어붙이기 → LfD 변환"],
            ["정책 평가", "E10: 시뮬 재생 + A/B 비교\n+ 전문가 태그",
             "정량 메트릭 + 주관 평가 병행"],
            ["안전 경계", "E11: Three.js 시각 편집기\n+ URDF joint limit 연동",
             "관절/토크 한계, 충돌 영역 시각 설정"],
        ], first_col_bold=True
    )

    # 3.6
    doc.add_heading("3.6  [P1~P7] 플랫폼 인프라", level=2)
    add_table(doc,
        ["항목", "기술 선정", "근거"],
        [
            ["프론트엔드", "Next.js 14+ (App Router)",
             "SSR, ISR, API Routes 통합"],
            ["백엔드", "FastAPI (Python)",
             "비동기, WebSocket 네이티브, 과학 생태계 직결"],
            ["인증", "NextAuth.js 또는 Keycloak",
             "OAuth2, SSO"],
            ["실시간 통신", "WebSocket (현재)\n→ WebTransport (2027~)",
             "99% 브라우저 지원"],
            ["GPU 스케줄링", "SLURM + Fair-share",
             "사용자별 공정 배분, 작업 큐, 우선순위"],
            ["적응형 품질", "클라이언트 벤치마크\n(접속 시 자동 측정)",
             "FPS, 네트워크 RTT, 메모리 → 품질 티어 자동 선택"],
            ["랜딩 페이지", "Next.js SSR + Three.js",
             "비로그인 로봇 3D 프리뷰, 기능 갤러리"],
            ["로봇 모델 라이브러리", "URDF 파일 + 메타데이터 DB\n+ USB VID/PID 자동 식별",
             "로봇 선택 → 시뮬/시각화 즉시 로드"],
        ], first_col_bold=True
    )

    # ================================================================
    # Part 4. 시스템 아키텍처
    # ================================================================
    doc.add_heading("4. 시스템 아키텍처", level=1)

    # 4.1
    doc.add_heading("4.1  전체 시스템 구성도", level=2)
    add_body(doc,
        "PhAI Studio는 로봇/모듈(Robot Side), 로컬 PC(User Side), "
        "서버(Cloud/On-Premise)의 3-Tier 아키텍처로 구성된다. "
        "사용자가 보고 조작하는 모든 화면은 100% 웹 브라우저이며, "
        "로컬 에이전트는 UI가 없는 백그라운드 서비스로 HW 통신만 담당한다."
    )
    add_body(doc, "[다이어그램 F — 추후 삽입: 통합 시스템 뷰]", bold=True)
    add_code_block(doc,
        "=== PhAI Studio 3-Tier Architecture ===\n"
        "\n"
        "[로봇/모듈]                    [로컬 PC]                      [서버]\n"
        "\n"
        " XM10 ─USB─┐                  ┌─ 로컬 에이전트 ──WebSocket──┐\n"
        " MD ──CAN──┤──USB CDC──┤      │  (PySerial, h5py)          │\n"
        " SM ──CAN──┘           │      │  - 1kHz 수집               │\n"
        "                       └──────┤  - HDF5 저장               │\n"
        "                              │  - FW Upload 프록시        │\n"
        "                              └──────────────┐              │\n"
        "                                             │              │\n"
        "                              ┌─ 브라우저 ◀──WebSocket──────┤\n"
        "                              │  - Monaco (코드 편집)      │\n"
        "                              │  - MuJoCo WASM (시뮬)     │\n"
        "                              │  - Three.js (3D URDF)     │\n"
        "                              │  - uPlot (그래프)          │\n"
        "                              │  - ONNX RT Web (AI 추론)  │\n"
        "                              └─────────────────────────────┤\n"
        "                                                           │\n"
        "                              ┌─ 서버 ─────────────────────┘\n"
        "                              │  - FastAPI (백엔드)\n"
        "                              │  - Docker GCC (클라우드 빌드)\n"
        "                              │  - MuJoCo MJX (GPU 학습)\n"
        "                              │  - SLURM (작업 스케줄링)\n"
        "                              │  - HDF5 데이터셋 DB\n"
        "                              └───────────────────"
    )

    # 4.2
    doc.add_heading("4.2  핵심 데이터 경로", level=2)
    add_table(doc,
        ["경로", "데이터", "주기", "프로토콜"],
        [
            ["XM10 → 에이전트", "센서/제어 데이터", "1kHz", "USB CDC (PhAI V2.2)"],
            ["에이전트 → HDF5", "전 채널 로우 데이터", "1kHz", "로컬 파일 I/O"],
            ["에이전트 → 브라우저", "다운샘플링 시각화 데이터", "60Hz", "WebSocket (localhost)"],
            ["브라우저 → 서버", "C 소스 코드, 학습 요청", "이벤트", "HTTP / WebSocket"],
            ["서버 → 브라우저", "빌드 결과, 학습 진행률", "이벤트/스트림", "WebSocket"],
            ["MuJoCo WASM (브라우저)", "시뮬 상태", "1kHz 계산, 60Hz 렌더", "로컬 (네트워크 무관)"],
        ], first_col_bold=True
    )

    # 4.3
    doc.add_heading("4.3  Connection-State UI Router", level=2)
    add_body(doc,
        "CTO 요구에 따라, 사용자의 로그인 상태와 로봇 연결 상태에 따라 UI가 자동으로 분기된다."
    )
    add_body(doc, "[다이어그램 G — 추후 삽입: Connection-State UI Router]", bold=True)

    add_body(doc, "State 1: 비로그인 (Not Authenticated)", bold=True)
    add_bullet(doc,
        "P6 Landing Page: PhAI Studio 소개, 데모 영상, 로봇 3D 프리뷰, 로그인 유도"
    )

    add_body(doc, "State 2: 로그인 + 로봇 Disconnected → Offline Dashboard", bold=True)
    add_bullet(doc, "저장 데이터 불러오기/시각화/저장 (D1~D6)")
    add_bullet(doc, "Robot Model Library에서 로봇 선택 → 동역학 시뮬레이션 (T1)")
    add_bullet(doc, "AI 학습 요청/모니터링 (T3~T4) — 서버 GPU, 로봇 연결 불필요")
    add_bullet(doc, "Web IDE 오프라인 코딩 (E1~E3) — 빌드는 서버에서")
    add_bullet(doc, "데이터 후처리/분석 (D4~D6)")

    add_body(doc, "State 3: 로그인 + 로봇 Connected → Live Dashboard", bold=True)
    add_bullet(doc, "접속 로봇 자동 식별 → 해당 URDF 3D 렌더링으로 환영 (P7 + D3)")
    add_bullet(doc, "실시간 1kHz 데이터 모니터링 (S1~S2)")
    add_bullet(doc, "데이터 수집/저장 (D1)")
    add_bullet(doc, "Live Digital Twin (E8)")
    add_bullet(doc, "FW Upload (E5)")
    add_bullet(doc, "모든 Offline 기능도 사용 가능")

    add_body(doc,
        "핵심 설계 원칙: Disconnected 상태에서도 Studio의 대부분 기능이 동작한다. "
        "로봇 연결은 실시간 데이터 수집과 FW 업로드에만 필수이다.",
        bold=True
    )

    # ================================================================
    # Part 5. 아키텍처 결정 사항
    # ================================================================
    doc.add_heading("5. 아키텍처 결정 사항", level=1)
    add_body(doc,
        "PhAI Studio 개발 과정에서 확정된 8가지 핵심 아키텍처 결정이다."
    )
    add_table(doc,
        ["#", "결정", "내용", "근거"],
        [
            ["1", "웹 프론트엔드 + 로컬 에이전트 Hybrid",
             "모든 UI는 웹 브라우저.\nHW 통신은 에이전트.\n'로컬 앱' 개념 없음.",
             "CTO Q1, Q2 대응.\n에이전트 = UI 없는 백그라운드 서비스\n(프린터 드라이버 비유)"],
            ["2", "에이전트 Primary + Web Serial Lite",
             "에이전트가 1kHz 무손실 + HDF5.\nWeb Serial은 에이전트 없는 과도기.",
             "PySerial: OS 네이티브 안정성.\nWeb Serial: Chrome 전용, 제약 있음"],
            ["3", "MuJoCo 2-Tier",
             "인터랙티브: WASM (브라우저)\n대규모 학습: MJX (서버 GPU)",
             "인터랙티브: 네트워크 지연 0\n학습: GPU 병렬 65만 step/s"],
            ["4", "C + Python 모두 제공",
             "C (XM_API): Sim-to-Real 동일 코드\nPython: 프로토타이핑, RL 스크립트",
             "C → WASM 컴파일로 브라우저 실행\nPython → 서버 MuJoCo API"],
            ["5", "URDF 표준",
             "로봇 모델 형식 통일",
             "업계 표준. MuJoCo + Three.js 양쪽 호환"],
            ["6", "V2.2 스트리밍 + AGR FTP 부트로더",
             "데이터: PhAI V2.2 (COBS+CRC16)\nFW: AGR FTP (Backup-First)",
             "검증됨. V3.0은 에이전트 단계에서 점진 도입"],
            ["7", "LfD (Learning from Demonstration)",
             "마우스 드래그 → 궤적 학습\n전문가 시범 → 모방 학습",
             "ONE PhAI Shared Autonomy 철학.\nACT, Diffusion Policy (LeRobot)"],
            ["8", "Connection-State UI Router",
             "비로그인/Disconnected/Connected\n3-State UI 분기",
             "CTO UX 흐름 대응.\nDisconnected에서도 대부분 기능 동작"],
        ], first_col_bold=True
    )

    # ================================================================
    # Part 6. Phase 로드맵
    # ================================================================
    doc.add_heading("6. Phase 로드맵", level=1)
    add_table(doc,
        ["Phase", "기간", "핵심 내용", "산출물"],
        [
            ["Phase 0\nFoundation", "3개월",
             "로컬 에이전트 MVP (PySerial + WebSocket + HDF5)\n"
             "uPlot 실시간 그래프\n"
             "FW Upload (Web Serial, AGR FTP)\n"
             "Docker 클라우드 빌드 파이프라인\n"
             "P6 Landing Page\n"
             "Connection-State UI Router",
             "에이전트 설치 파일\n웹 대시보드 MVP\nDocker 빌드 이미지\n랜딩 페이지"],
            ["Phase 1\nCore Platform", "6개월",
             "Monaco Web IDE (Sandbox + Direct Coding)\n"
             "Three.js URDF 3D 시각화\n"
             "MuJoCo WASM 인터랙티브 시뮬레이터\n"
             "데이터 후처리 (SciPy)\n"
             "XM_API Simulation Stubs (C → WASM)\n"
             "Live Digital Twin (E8)\n"
             "Safety Boundary Editor (E11)\n"
             "P7 Robot Model Library",
             "Web IDE 정식 버전\n시뮬레이터 정식 버전\nPhAI-X1 URDF 모델"],
            ["Phase 2\nAI & Simulation", "6개월",
             "MuJoCo MJX GPU 학습 서버 (SLURM)\n"
             "RL/LfD 학습 파이프라인\n"
             "Sim-to-Real (ONNX → STM32Cube.AI)\n"
             "Domain Randomization\n"
             "Expert Demo Recorder (E7)\n"
             "Trajectory Replay (E9)\n"
             "Policy Eval Dashboard (E10)\n"
             "DAgger 점진적 학습\n"
             "LLM-Assisted Coding",
             "GPU 학습 서버\nSim-to-Real 파이프라인\nLfD 도구 체인"],
            ["Phase 3\nActuator", "3개월",
             "SAM_GUI → Python/Web 마이그레이션\n"
             "System Identification (SciPy)\n"
             "Actuator Parameter DB\n"
             "Sim-to-Real 자동 파라미터 적용",
             "시스템 식별 웹 도구\nActuator DB"],
            ["Phase 4\nScale", "지속",
             "다중 로봇 URDF 추가 (4족/매니퓰/휴머노이드)\n"
             "다수 사용자 최적화\n"
             "적응형 품질 시스템\n"
             "전문가 데이터 마켓플레이스",
             "다중 로봇 지원\n사용자 확장"],
        ], first_col_bold=True
    )


# ============================================================
# English Version (STUB)
# ============================================================
def build_phai_studio_spec_en(doc):
    """PhAI Studio Integrated Technical Specification (English) — Stub."""

    add_title_page(doc,
        "PhAI Studio\nIntegrated Technical Specification",
        "Physical AI Platform — π0 Dual-System + Human Expert Framework",
        "v1.0  ·  2026-03-12  ·  Angel Robotics"
    )

    add_body(doc,
        "English version is under preparation. "
        "Please refer to the Korean version for the complete specification."
    )


def main():
    doc_kr = setup_document(FONT_KR)
    build_phai_studio_spec_kr(doc_kr)
    kr_path = os.path.join(OUTPUT_DIR, "PhAI_Studio_통합_기술사양서_v1.0.docx")
    doc_kr.save(kr_path)
    print(f"[OK] Korean: {kr_path}")

    doc_en = setup_document(FONT_KR)
    build_phai_studio_spec_en(doc_en)
    en_path = os.path.join(OUTPUT_DIR, "PhAI_Studio_Integrated_Tech_Spec_v1.0.docx")
    doc_en.save(en_path)
    print(f"[OK] English: {en_path}")


if __name__ == "__main__":
    main()
