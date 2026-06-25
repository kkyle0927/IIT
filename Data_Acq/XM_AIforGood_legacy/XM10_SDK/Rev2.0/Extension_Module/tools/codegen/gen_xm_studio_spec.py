"""
XM-Studio Web IDE 기술기획서 Word(.docx) 생성기.

XM10 임베디드 시스템을 위한 웹 기반 개발 환경(XM-Studio) 구축
기술 사양 확정 및 개발 기획 문서.

Usage:
    python docs/xm-studio/gen_xm_studio_spec.py

Dependencies:
    pip install python-docx pyyaml
"""
import sys
import os

BASE_DIR = r"C:\AGR_HyundoKim\02_PhAI\PhAI_Web\Docs\Plan\PhAI_Studio_기획"
sys.path.insert(0, BASE_DIR)
sys.stdout.reconfigure(encoding="utf-8")

from gen_docx_v2_base import (
    setup_document, add_title_page, add_info_cards, add_table, add_code_block,
    add_body, add_bullet, add_ref, FONT_KR,
    BG_INFO_BLUE, BG_INFO_GRAY, BG_INFO_GREEN, BG_INFO_AMBER,
)

F = FONT_KR
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def _add_toc(doc, items, lang="kr"):
    from docx.shared import Pt
    from gen_docx_v2_base import CLR_GRAY
    doc.add_heading("목 차" if lang == "kr" else "Table of Contents", level=1)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("   " + item)
        r.font.name = F
        r.font.size = Pt(10)
        r.font.color.rgb = CLR_GRAY
    doc.add_page_break()


# ============================================================
# Korean Version
# ============================================================
def build_xm_studio_spec_kr(doc):
    """XM-Studio Web IDE 기술기획서 (한국어)."""

    add_title_page(doc,
        "XM-Studio Web IDE\n기술 기획서",
        "XM10 임베디드 시스템 웹 기반 개발 환경 — 기술 사양 및 개발 기획",
        "v1.0  ·  2026-03-11  ·  Angel Robotics"
    )

    add_info_cards(doc, [
        ("적용 대상", "PhAI Studio Web + XM10 FW", BG_INFO_BLUE),
        ("핵심 원칙", "XM_API 제약 기반 IDE", BG_INFO_GRAY),
        ("빌드 방식", "Docker + libXM_Full.a", BG_INFO_GREEN),
        ("FW 전송", "Web Serial FTP (서버 미경유)", BG_INFO_GREEN),
    ])
    add_info_cards(doc, [
        ("개발 우선순위", "Sandbox + 직접 코딩 우선", BG_INFO_BLUE),
        ("AI LLM", "기획 단계 (구현 미착수)", BG_INFO_AMBER),
        ("버전", "v1.0 (기술 기획서)", BG_INFO_GRAY),
        ("상태", "기획 확정, 개발 준비 중", BG_INFO_GRAY),
    ])

    toc_items = [
        "1. 개요",
        "2. 시스템 아키텍처",
        "3. Phase 0: 인프라 기반 구축 사양",
        "4. Phase 1: 핵심 기능 사양 (Sandbox + 직접 코딩)",
        "5. Phase 2: AI LLM 코드 생성 (기획 단계)",
        "6. 개발 우선순위 및 일정",
        "7. 담당 분야 (업무 배분 기준)",
        "8. 리스크 및 제약 조건",
        "9. 레퍼런스",
    ]
    _add_toc(doc, toc_items, "kr")

    # ================================================================
    # Section 1
    # ================================================================
    doc.add_heading("1. 개요", level=1)

    doc.add_heading("1.1  목적", level=2)
    add_body(doc,
        "본 문서는 XM10 임베디드 시스템을 위한 웹 기반 개발 환경(XM-Studio)의 "
        "기술 사양과 개발 기획을 정의한다. STM32CubeIDE 없이 웹 브라우저만으로 "
        "XM10 사용자 코드를 작성, 컴파일, 업로드할 수 있는 통합 개발 플랫폼을 구축한다."
    )
    add_body(doc,
        "로블록스(Roblox) 스튜디오 패러다임을 채택하여, 사용자는 비즈니스 로직(제어 알고리즘, "
        "데이터 수집 등)에만 집중하고 인프라(HAL, RTOS, 드라이버)는 플랫폼이 제공하는 구조이다."
    )

    doc.add_heading("1.2  적용 범위", level=2)
    add_bullet(doc, "PhAI Studio Frontend (Next.js + TypeScript + Monaco Editor)")
    add_bullet(doc, "Cloud Compile Service (FastAPI + Docker + ARM GCC)")
    add_bullet(doc, "XM10 SDK 기반 Static Library (libXM_Full.a)")
    add_bullet(doc, "FW Upload (Web Serial FTP — FW Upload 기술사양서 v1.1 연동)")
    add_bullet(doc, "AI LLM 코드 생성 (기획 단계, Phase 2)")

    doc.add_heading("1.3  핵심 설계 원칙", level=2)
    add_table(doc,
        ["원칙", "설명"],
        [
            ["XM_API Facade 제약", "사용자는 XM_API 함수만 호출 가능. HAL, RTOS, 내부 드라이버 직접 호출 금지"],
            ["서버 미경유 직접 전송", "FW 바이너리는 Browser → USB CDC로 직접 전송 (FW Upload Spec v1.1)"],
            ["Static Library 빌드", "단일 libXM_Full.a + 사용자 코드만 컴파일하여 빌드 시간 최소화"],
            ["개발 우선순위", "Sandbox + 직접 코딩 먼저 개발, AI LLM은 서버 인프라 구축 후 착수"],
            ["packaged.bin 생성", "클라우드 빌드가 AGR_FwInfo_t 헤더 포함 packaged.bin까지 생성"],
        ], first_col_bold=True
    )

    doc.add_heading("1.4  참조 문서", level=2)
    add_table(doc,
        ["문서", "위치", "설명"],
        [
            ["PhAI Studio FW Upload\n기술사양서 v1.1", "docs/bootloader/", "FTP 프로토콜, AGR_FwInfo_t, 업로드 워크플로우"],
            ["XM V2 Bootloader Plan", "plan 파일", "BL 전체 설계 (Flash Layout, Core, Transport)"],
            ["Total Data Packet 설계사양서", "docs/total_data_packet/", "Module ID 체계, USB CDC 스트리밍"],
            ["XM10 SDK 릴리즈 레포", "Extension_Module/", ".a 파일, XM_API, .cproject"],
        ], first_col_bold=True
    )

    doc.add_heading("1.5  용어 정의", level=2)
    add_table(doc,
        ["용어", "정의"],
        [
            ["XM_API", "XM10 사용자 Facade API (~80개 함수). User_Setup/User_Loop에서 호출"],
            ["libXM_Full.a", "Core + HAL + CMSIS + FreeRTOS + Middleware + XM_FW를 단일 통합한 Static Library"],
            ["packaged.bin", "AGR_FwInfo_t 헤더(1KB) + App Binary(max 767KB). 부트로더가 수신하는 형식"],
            ["AGR FTP", "Angel Robotics 독자 FW Transfer Protocol over USB CDC (PhAI V2.2 패킷)"],
            ["Sandbox", "사전 구성된 예제 코드를 로드하여 즉시 수정/빌드할 수 있는 갤러리"],
            ["Vibe Coding", "자연어 프롬프트로 LLM이 코드를 생성하는 개발 패러다임"],
        ], first_col_bold=True
    )

    # ================================================================
    # Section 2
    # ================================================================
    doc.add_heading("2. 시스템 아키텍처", level=1)

    doc.add_heading("2.1  전체 구성도", level=2)
    add_body(doc,
        "XM-Studio는 코드 편집(브라우저) → 클라우드 컴파일(서버) → FW 업로드(브라우저 → USB 직접 전송)의 "
        "3단계 파이프라인으로 구성된다. 컴파일은 서버를 경유하지만, FW 전송은 서버를 경유하지 않는다."
    )
    add_code_block(doc,
        "=== XM-Studio 전체 아키텍처 ===\n"
        "\n"
        "  [사용자 브라우저]\n"
        "       | Monaco Editor (C 코드 편집)\n"
        "       | Sandbox (예제 로드/수정)\n"
        "       v\n"
        "  [PhAI Studio Frontend (Next.js)]\n"
        "       | POST /api/compile (다중 파일 소스 전송)\n"
        "       v\n"
        "  [Cloud Compile Server (FastAPI + Docker)]\n"
        "       | Docker Container: arm-none-eabi-gcc\n"
        "       | CMake + libXM_Full.a 링크\n"
        "       | objcopy -> raw .bin -> agr_fw_packager -> packaged.bin\n"
        "       v\n"
        "  [packaged.bin 다운로드 -> 브라우저 메모리]\n"
        "       | FTP Host Engine (TypeScript)\n"
        "       v\n"
        "  [Web Serial API] ---- USB CDC, 921600 baud\n"
        "       |\n"
        "  ======== USB Cable (물리 연결) ========\n"
        "       |\n"
        "  [XM10 Bootloader -> Flash 기록]\n"
        "\n"
        "  핵심: 컴파일은 서버 경유 / FW 전송은 서버 미경유 (직접 USB)"
    )

    doc.add_heading("2.2  핵심 설계 원칙 상세", level=2)
    add_body(doc, "XM_API 제약 기반 IDE (API-Constrained IDE):", bold=True)
    add_bullet(doc, "사용자 코드는 XM_API 헤더만 include 가능 — HAL, RTOS, 내부 드라이버 헤더 include 불가")
    add_bullet(doc, "User_Setup() / User_Loop() 두 함수가 진입점 — 반드시 정의 필수 (빌드 전 사전 검증)")
    add_bullet(doc, "XM_Apps 폴더 범위 내에서 다중 파일/폴더 자유 생성 가능")
    add_body(doc, "로블록스 패러다임:", bold=True)
    add_bullet(doc, "사용자는 게임 로직(비즈니스 로직)만 작성 — 엔진(HAL+RTOS+드라이버)은 플랫폼이 제공")
    add_bullet(doc, "libXM_Full.a = '게임 엔진', 사용자 코드 = '게임 스크립트'")

    doc.add_heading("2.3  레퍼런스 플랫폼 비교", level=2)
    add_body(doc,
        "유사 임베디드 웹 개발 플랫폼과의 기능 비교:",
        bold=True
    )
    add_table(doc,
        ["플랫폼", "에디터", "빌드", "업로드", "AI", "XM-Studio 차별점"],
        [
            ["Arduino Web Editor", "Monaco", "클라우드\n(Pre-compiled Core)", "Web Serial\n(esptool)", "Copilot\n(일부)", "API 제약 없음\n(전체 HAL 노출)"],
            ["Keil Studio Cloud", "VS Code Web", "클라우드\n(ARMC6)", "CMSIS-DAP", "없음", "ARM 전용\n브라우저 제한"],
            ["PlatformIO Remote", "VS Code", "원격 Agent", "Agent 경유", "없음", "Agent 설치 필요"],
            ["Wokwi", "Monaco", "로컬/CI", "시뮬레이터", "없음", "시뮬레이터 전용\n(실물 업로드 없음)"],
            ["XM-Studio (본 설계)", "Monaco", "Docker +\nlibXM_Full.a", "Web Serial FTP\n(packaged.bin)", "RAG\n(Phase 2 계획)", "XM_API 제약 IDE\npackaged.bin FTP\nBackup-First 안전"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] Arduino Cloud Builder — blog.arduino.cc/2024/12/13")
    add_ref(doc, "[Ref] Keil Studio Cloud — keil.arm.com")
    add_ref(doc, "[Ref] ESP-IDF Docker Image — docs.espressif.com")
    add_ref(doc, "[Ref] Wokwi STM32 Simulator — wokwi.com/stm32")

    # ================================================================
    # Section 3
    # ================================================================
    doc.add_heading("3. Phase 0: 인프라 기반 구축 사양", level=1)

    doc.add_heading("3.1  libXM_Full.a 단일 통합 Static Library", level=2)
    add_body(doc,
        "XM10 FW의 모든 비사용자 코드를 단일 .a 파일로 통합하여, "
        "클라우드 빌드 시 사용자 코드만 컴파일 + 링크하는 전략이다. "
        "Arduino Cloud가 Pre-compiled Core로 빌드 시간을 50% 단축한 것과 동일한 접근법이다."
    )
    add_body(doc, "포함 대상 (~220개 .o):", bold=True)
    add_table(doc,
        ["계층", "포함 내용", "예상 .o 수"],
        [
            ["Core", "core_process.c, system_init.c 등", "~10"],
            ["HAL", "STM32H7 HAL 드라이버", "~40"],
            ["CMSIS", "system_stm32h7xx.c, DSP 라이브러리", "~15"],
            ["FreeRTOS", "RTOS 커널 + CMSIS-OS2 래퍼", "~20"],
            ["Middleware", "FatFs, USB Host/Device", "~30"],
            ["Compatible", "stm32h7xx_compatible.c 등", "~5"],
            ["XM_FW", "XM_API 구현, IOIF, Devices, Services", "~100"],
        ], first_col_bold=True
    )
    add_body(doc, "제외 대상:", bold=True)
    add_bullet(doc, "startup ASM (startup_stm32h743xihx.s) — Docker에 별도 복사, 매 빌드 시 어셈블")
    add_bullet(doc, "사용자 코드 (XM_Apps/) — 사용자가 웹에서 작성하는 대상")
    add_body(doc, "컴파일 및 링크 전략:", bold=True)
    add_table(doc,
        ["항목", "설정", "이유"],
        [
            ["컴파일 옵션", "-ffunction-sections\n-fdata-sections", "함수/데이터 단위 섹션 분리 → 미사용 코드 제거 가능"],
            ["링커: --whole-archive", "libXM_Full.a에 적용", "core_process.c의 extern User_Setup/User_Loop\n심볼이 사용자 코드에서 해결되도록 강제 포함"],
            ["링커: --gc-sections", "최종 링크에 적용", "--whole-archive로 포함된 .o 중\n실제 미사용 함수/데이터 제거 → 바이너리 최적화"],
            ["Toolchain 고정", "arm-none-eabi-gcc\n12.3.rel1", ".a 생성과 사용자 코드 컴파일에 동일 버전 사용 필수\n(Mbed OS toolchain 호환성 이슈 방지)"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] Arduino Cloud: pre-compiled core로 빌드 50% 단축 — blog.arduino.cc/2024/12/13")
    add_ref(doc, "[Ref] Mbed OS: .a toolchain 호환성 이슈 — os.mbed.com/questions/83007")

    doc.add_heading("3.2  Docker SDK 패키지", level=2)
    add_body(doc,
        "클라우드 빌드에 필요한 모든 파일을 Docker 이미지에 탑재한다. "
        "ESP-IDF의 espressif/idf Docker 이미지 패턴을 참고하여 "
        "버전 고정된 빌드 환경을 제공한다."
    )
    add_code_block(doc,
        "Docker SDK 패키지 디렉토리 레이아웃:\n"
        "\n"
        "  /opt/xm-sdk/\n"
        "  ├── lib/\n"
        "  │   └── libXM_Full.a          # 단일 통합 Static Library\n"
        "  ├── include/\n"
        "  │   ├── XM_API/               # xm_api.h + 하위 헤더\n"
        "  │   ├── internal/             # XM_API 전이 의존성 헤더 (~10개)\n"
        "  │   │   ├── cm_drv.h\n"
        "  │   │   ├── data_object*.h\n"
        "  │   │   ├── task_mngr.h\n"
        "  │   │   └── module.h 등\n"
        "  │   └── rtos/                 # cmsis_os2.h (XM_API 전이 의존)\n"
        "  ├── linker/\n"
        "  │   └── STM32H743XIHx_FLASH.ld\n"
        "  ├── startup/\n"
        "  │   └── startup_stm32h743xihx.s  # 매 빌드 시 어셈블\n"
        "  ├── cmake/\n"
        "  │   └── cloud_build.cmake     # 사용자 코드 자동 탐색 + 링크\n"
        "  └── tools/\n"
        "      └── agr_fw_packager.py    # raw .bin -> packaged.bin 변환"
    )
    add_body(doc, "Dockerfile 핵심 구성:", bold=True)
    add_table(doc,
        ["항목", "값", "비고"],
        [
            ["Base Image", "ubuntu:22.04", "LTS 안정성"],
            ["Toolchain", "arm-none-eabi-gcc 12.3.rel1", "libXM_Full.a와 동일 버전 필수"],
            ["Build System", "CMake 3.28 + Ninja", "빠른 빌드"],
            ["Python", "3.10+", "agr_fw_packager.py 실행용"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] ESP-IDF Docker Image — docs.espressif.com/projects/esp-idf/en/latest/esp32/api-guides/tools/idf-docker-image.html")

    doc.add_heading("3.3  XM_API JSON 스키마", level=2)
    add_body(doc,
        "XM_API 함수 시그니처, 파라미터, 반환값, 타입 정의를 JSON으로 구조화하여 "
        "Monaco Editor 자동완성과 향후 AI RAG 데이터로 활용한다."
    )
    add_table(doc,
        ["항목", "내용"],
        [
            ["대상 함수", "~80개 XM_API public 함수"],
            ["포함 타입", "enum, struct, typedef (XM_API에서 노출하는 것만)"],
            ["활용처", "Monaco registerCompletionItemProvider (Phase 1)\nAI RAG 데이터 (Phase 2)"],
            ["생성 방식", "XM_API 헤더 파싱 스크립트 (Python + pycparser 또는 수동)"],
        ], first_col_bold=True
    )
    add_code_block(doc,
        "JSON 스키마 예시:\n"
        "{\n"
        '  "functions": [\n'
        "    {\n"
        '      "name": "XM_GPIO_Write",\n'
        '      "brief": "GPIO 핀 출력 설정",\n'
        '      "params": [\n'
        '        {"name": "pin", "type": "XM_GPIO_Pin_t", "desc": "GPIO 핀 번호"},\n'
        '        {"name": "state", "type": "XM_GPIO_State_t", "desc": "HIGH/LOW"}\n'
        "      ],\n"
        '      "return": {"type": "void"},\n'
        '      "category": "GPIO"\n'
        "    }\n"
        "  ],\n"
        '  "types": [\n'
        "    {\n"
        '      "name": "XM_GPIO_Pin_t",\n'
        '      "kind": "enum",\n'
        '      "values": ["XM_GPIO_PIN_0", "XM_GPIO_PIN_1", ...]\n'
        "    }\n"
        "  ]\n"
        "}"
    )

    doc.add_heading("3.4  packaged.bin 빌드 파이프라인", level=2)
    add_body(doc,
        "클라우드 빌드의 최종 산출물은 packaged.bin이다. "
        "부트로더의 FTP 프로토콜이 이 형식을 기대하며, "
        "브라우저의 FTP Host Engine이 AGR_FwInfo_t 헤더를 파싱하여 호환성을 사전 검증한다."
    )
    add_code_block(doc,
        "빌드 파이프라인:\n"
        "\n"
        "  [사용자 코드 (*.c, *.h)]\n"
        "       | arm-none-eabi-gcc -c (컴파일)\n"
        "       v\n"
        "  [사용자 .o 파일들]\n"
        "       | arm-none-eabi-gcc (링크)\n"
        "       | --whole-archive libXM_Full.a --no-whole-archive\n"
        "       | --gc-sections -T STM32H743XIHx_FLASH.ld\n"
        "       | startup_stm32h743xihx.o\n"
        "       v\n"
        "  [firmware.elf]\n"
        "       | arm-none-eabi-objcopy -O binary\n"
        "       v\n"
        "  [firmware.bin] (raw binary)\n"
        "       | agr_fw_packager.py\n"
        "       | --fw_version=X.Y.Z --target=XM --hw_rev_min=0x11 --hw_rev_max=0x20\n"
        "       v\n"
        "  [packaged.bin] = AGR_FwInfo_t(1KB) + firmware.bin(max 767KB)\n"
        "       |\n"
        "       v\n"
        "  브라우저로 다운로드 -> FTP Host Engine -> Web Serial -> XM BL"
    )
    add_body(doc,
        "packaged.bin 구조 (FW Upload 기술사양서 v1.1 Section 5 참조):",
        bold=True
    )
    add_table(doc,
        ["영역", "크기", "내용"],
        [
            ["AGR_FwInfo_t 헤더", "1KB (64B 유효 + 960B 패딩)", "signature, fw_size, fw_crc32, fw_version,\nbuild_timestamp, target_module,\nhw_rev_min, hw_rev_max, flags, header_crc16"],
            ["App Binary", "max 767KB", "arm-none-eabi-objcopy로 추출한 raw binary"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] PhAI Studio FW Upload 기술사양서 v1.1 — Section 5. AGR_FwInfo_t 구조체")

    # ================================================================
    # Section 4
    # ================================================================
    doc.add_heading("4. Phase 1: 핵심 기능 사양 (Sandbox + 직접 코딩)", level=1)
    add_body(doc,
        "Phase 1 = Sandbox + 직접 타이핑 + 클라우드 빌드 + FW 업로드. AI LLM은 Phase 2로 분리.",
        bold=True
    )

    doc.add_heading("4.1  Monaco Editor C 에디터", level=2)
    add_body(doc,
        "VS Code의 코어 에디터인 Monaco Editor를 사용하여 C 언어 편집 환경을 제공한다."
    )
    add_table(doc,
        ["기능", "구현 방식", "비고"],
        [
            ["구문 강조", "Monaco 내장 C 언어 지원", "기본 제공"],
            ["XM_API 자동완성", "registerCompletionItemProvider\n+ XM_API JSON 스키마", "커스텀 구현 필수\n(Monaco는 C LSP 미내장)"],
            ["컴파일 에러 표시", "빌드 서버 GCC 출력 파싱\n→ Monaco setModelMarkers", "서버 의존\n(클라이언트 실시간 검증 불가)"],
            ["다중 파일/폴더", "파일 탐색기 (트리뷰)\n다중 탭, 새 파일/폴더 생성", "커스텀 UI 구현"],
            ["User_Setup/Loop 검증", "빌드 전 클라이언트 사전 검사\n(정규식 매칭)", "빌드 실패 사전 방지"],
            ["괄호 매칭/포맷팅", "Monaco 내장", "기본 제공"],
        ], first_col_bold=True
    )
    add_body(doc, "Monaco C 언어 한계 사항:", bold=True)
    add_bullet(doc, "Monaco는 C 언어에 대해 syntax highlighting만 네이티브 지원하며, 실시간 에러 검증은 제공하지 않음")
    add_bullet(doc, "에러 검증은 클라우드 빌드 서버의 GCC 출력에 의존 (Save → Build → 에러 표시 사이클)")
    add_bullet(doc, "향후 확장: Clangd LSP를 WebSocket 경유로 연동하면 실시간 에러 표시 가능 (Phase 2+ 검토)")
    add_ref(doc, "[Ref] Monaco Editor 통합 가이드 — blog.spectralcore.com/integrating-monaco-editor")
    add_ref(doc, "[Ref] Monaco + LSP 설정 — dev.to/__4f1641/so-you-want-to-set-up-a-monaco-editor")

    doc.add_heading("4.2  Sandbox (예제 갤러리)", level=2)
    add_body(doc,
        "기존 XM10 예제 코드를 카테고리별로 정리하여, 사용자가 예제를 선택 → 에디터 로드 → "
        "자유 수정 → 빌드 → FW 업로드하는 워크플로우를 제공한다."
    )
    add_table(doc,
        ["항목", "내용"],
        [
            ["예제 소스", "52개 기존 예제\n(ARC_ExtensionBoard/.../Examples/)"],
            ["카테고리", "입문(6), 초급(12), 중급(15), 고급(19)"],
            ["메타데이터", "JSON: id, title, category, difficulty,\ndescription, files[], required_hw"],
            ["워크플로우", "예제 선택 → Editor 로드 → 수정 → 빌드 → 업로드"],
            ["예제별 README", "사용 API 목록, 동작 설명, HW 연결 안내"],
        ], first_col_bold=True
    )
    add_body(doc,
        "CodeSandbox 대비 차이점: XM-Studio의 Sandbox는 클라우드 VM이 아니라 "
        "'정적 예제 코드 템플릿 로딩 → Monaco에서 편집 → 클라우드 빌드'로, "
        "훨씬 가벼운 구조이다.",
        bold=True
    )
    add_ref(doc, "[Ref] CodeSandbox 템플릿 관리 — codesandbox.io/docs/sdk/templates")

    doc.add_heading("4.3  클라우드 컴파일 서비스", level=2)
    add_body(doc, "REST API 설계:", bold=True)
    add_table(doc,
        ["Method", "Endpoint", "설명"],
        [
            ["POST", "/api/compile", "다중 파일 소스 전송 → packaged.bin 반환"],
            ["GET", "/api/compile/{job_id}/status", "빌드 상태 조회 (pending/building/done/error)"],
            ["GET", "/api/compile/{job_id}/binary", "packaged.bin 다운로드"],
            ["GET", "/api/compile/{job_id}/errors", "GCC 에러/경고 (JSON 형식)"],
        ], first_col_bold=True
    )
    add_body(doc, "요청 형식:", bold=True)
    add_code_block(doc,
        "POST /api/compile\n"
        "{\n"
        '  "fw_version": "1.0.0",\n'
        '  "hw_rev_min": "0x11",\n'
        '  "hw_rev_max": "0x20",\n'
        '  "files": [\n'
        '    { "path": "user_apps.c", "content": "..." },\n'
        '    { "path": "my_algo.c", "content": "..." },\n'
        '    { "path": "my_algo.h", "content": "..." }\n'
        "  ]\n"
        "}"
    )
    add_body(doc, "빌드 파이프라인 상세:", bold=True)
    add_table(doc,
        ["단계", "동작", "예상 시간"],
        [
            ["1. 소스 수신", "JSON → 임시 디렉토리에 파일 생성", "< 0.1초"],
            ["2. CMake 설정", "cloud_build.cmake로 자동 설정", "< 0.5초"],
            ["3. 컴파일", "arm-none-eabi-gcc -c (사용자 코드만)", "1-3초"],
            ["4. 링크", "libXM_Full.a + 사용자 .o + startup.o", "1-2초"],
            ["5. objcopy", ".elf → raw .bin 변환", "< 0.1초"],
            ["6. 패키징", "agr_fw_packager.py → packaged.bin", "< 0.1초"],
            ["7. 응답", "packaged.bin + 메모리 사용량 보고", "< 0.1초"],
        ], first_col_bold=True
    )
    add_body(doc, "총 빌드 시간 목표: 3~8초", bold=True)
    add_body(doc, "에러 파싱 형식:", bold=True)
    add_code_block(doc,
        "GCC 출력 → JSON 변환:\n"
        "{\n"
        '  "errors": [\n'
        "    {\n"
        '      "file": "user_apps.c",\n'
        '      "line": 42,\n'
        '      "column": 15,\n'
        '      "severity": "error",\n'
        '      "message": "implicit declaration of function \'HAL_GPIO_WritePin\'"\n'
        "    }\n"
        "  ],\n"
        '  "warnings": [...],\n'
        '  "memory": {\n'
        '    "flash_used": 245760,\n'
        '    "flash_total": 786432,\n'
        '    "ram_used": 32768,\n'
        '    "ram_total": 524288\n'
        "  }\n"
        "}"
    )
    add_body(doc, "사용자 프로젝트 제약:", bold=True)
    add_table(doc,
        ["제약", "값", "이유"],
        [
            ["최대 파일 수", "20개", "Docker 빌드 시간 제한"],
            ["최대 총 크기", "500KB", "소스 코드 합산"],
            ["필수 함수", "User_Setup(), User_Loop()", "libXM_Full.a의 extern 참조 해결"],
            ["허용 include", "xm_api.h (+ 전이 의존)", "XM_API 외 헤더 include 금지"],
            ["빌드 타임아웃", "30초", "무한 루프/대용량 코드 방어"],
        ], first_col_bold=True
    )

    doc.add_heading("4.4  FW 업로드 연동 (Web Serial FTP)", level=2)
    add_body(doc,
        "FW 업로드 프로토콜의 상세 사양은 별도 문서(PhAI Studio FW Upload 기술사양서 v1.1)에 "
        "정의되어 있으며, 본 섹션은 XM-Studio 관점의 연동 사양만 기술한다.",
        bold=True
    )
    add_body(doc, "XM-Studio ↔ FW Upload 연동:", bold=True)
    add_table(doc,
        ["항목", "내용"],
        [
            ["빌드 산출물", "packaged.bin (클라우드 빌드가 AGR_FwInfo_t 헤더 포함하여 생성)"],
            ["전송 경로", "브라우저 메모리 → FTP Host Engine → Web Serial → USB CDC → XM BL"],
            ["서버 미경유", "packaged.bin은 브라우저에서 직접 USB 전송 (FW Upload Spec v1.1 원칙)"],
            ["호환성 검증", "AGR_FwInfo_t 헤더의 target_module, hw_rev_min/max로 사전 검증"],
            ["업로드 5단계", "Backup → Erase → Transfer → Verify → Commit"],
            ["HTTPS 필수", "Web Serial API는 보안 컨텍스트(HTTPS 또는 localhost)에서만 동작"],
        ], first_col_bold=True
    )
    add_body(doc, "FW 관리 확장성 로드맵 (FW Upload Spec v1.1 Section 13):", bold=True)
    add_table(doc,
        ["Phase", "기능", "설명"],
        [
            ["Phase A (현재)", "Custom Upload", "클라우드 빌드 결과를 브라우저에서 직접 USB 전송"],
            ["Phase B", "Official 버전 관리", "서버에 검증된 FW 저장, 버전 목록 제공"],
            ["Phase C", "OTA 배포", "서버 다운로드 → USB 전송, 호환 FW 자동 추천"],
            ["Phase D", "Fleet 관리", "다수 디바이스 일괄 업데이트, 배포 대시보드"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] PhAI Studio FW Upload 기술사양서 v1.1 — docs/bootloader/")
    add_ref(doc, "[Ref] MDN Web Serial API — developer.mozilla.org/en-US/docs/Web/API/Web_Serial_API")
    add_ref(doc, "[Ref] Bleskomat Web Serial FW Upload — github.com/bleskomat/bleskomat-web-serial")

    # ================================================================
    # Section 5
    # ================================================================
    doc.add_heading("5. Phase 2: AI LLM 코드 생성 (기획 단계)", level=1)
    add_body(doc,
        "Phase 1 개발 및 서버 인프라 구축 기간 동안 기획을 심화하는 단계이다. "
        "구현은 Phase 1 완료 후 별도 AI 기술사양서를 작성한 뒤 착수한다.",
        bold=True
    )

    doc.add_heading("5.1  기술 방향성 비교 검토", level=2)
    add_table(doc,
        ["방법", "구조", "장점", "단점", "비용"],
        [
            ["A: 외부 LLM API\n(Claude/GPT/Gemini)", "사용자 → Backend\n→ 외부 API\n→ 코드 생성", "최신 모델, 고성능\n빠른 도입", "API 비용(토큰 과금)\n외부 의존, 네트워크 필수\n사용자 코드 외부 전송", "종량제\n(요청당 과금)"],
            ["B: 자체 호스팅\n오픈소스 모델", "사용자 → Backend\n→ 로컬 GPU\n→ 코드 생성", "비용 고정(GPU 서버)\n데이터 외부 미전송\n오프라인 가능", "GPU 서버 필요\n모델 성능 한계\n유지보수 부담", "고정비\n(GPU 서버)"],
            ["C: 하이브리드", "간단 코드 = 로컬\n복잡 코드 = 외부", "비용/성능 균형", "구현 복잡도 증가\n라우팅 로직 필요", "혼합"],
            ["D: 사용자 API Key", "사용자가 자신의\nAPI Key 입력", "Angel Robotics 비용 0\n사용자 선택권", "UX 허들 높음\n설정 복잡\n사용자 비용 부담", "사용자 부담"],
        ], first_col_bold=True
    )
    add_body(doc, "현재 결론: 방법론 미확정. Phase 1 기간 중 비교 검토 진행, 서버 GPU 가용성 확인 후 최종 결정.", bold=True)

    doc.add_heading("5.2  Hardware Hallucination 리스크", level=2)
    add_body(doc,
        "임베디드 AI 코드 생성의 핵심 리스크는 'Hardware Hallucination'이다. "
        "LLM이 존재하지 않는 API, 잘못된 핀 번호, 부적절한 페리페럴 초기화 시퀀스를 "
        "생성하는 문제로, 일반 소프트웨어 코드 생성보다 위험도가 높다."
    )
    add_table(doc,
        ["리스크", "설명", "XM-Studio 대응"],
        [
            ["존재하지 않는 API 호출", "LLM이 HAL_GPIO_WritePin 등\nXM_API에 없는 함수 생성", "API 화이트리스트 검증\n(80개 XM_API 함수만 허용)"],
            ["잘못된 파라미터", "유효하지 않는 핀 번호,\n범위 밖 값 생성", "파라미터 범위 검증\n+ 컴파일 테스트"],
            ["타이밍 제약 미인식", "실시간 제어 루프에서\nblocking 함수 호출", "금지 패턴 검출\n(malloc, __disable_irq 등)"],
            ["구조 위반", "User_Setup/User_Loop 밖에\n전역 코드 배치", "구조 검증\n(필수 함수 존재 확인)"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] AI Code Generation in Embedded Systems — wedolow.com/resources/vibe-coding-ai-code-generation-embedded-systems")
    add_ref(doc, "[Ref] Embedded AI Constraints — gocodeo.com/post/ai-code-generation-in-embedded-systems-constraints-and-solutions")

    doc.add_heading("5.3  RAG vs Fine-Tuning vs Long Context 비교", level=2)
    add_table(doc,
        ["방식", "원리", "장점", "단점", "XM-Studio 적합도"],
        [
            ["RAG", "쿼리 시 관련 데이터\n검색 → 프롬프트 주입", "모델 변경 없음\n데이터 업데이트 용이\n비용 효율적", "검색 파이프라인 필요\n검색 품질에 의존", "높음\n(80개 API + 52예제\n= 적절한 규모)"],
            ["Fine-Tuning", "도메인 데이터로\n모델 가중치 조정", "도메인 특화 성능\n추론 시 추가 비용 없음", "학습 비용 높음\n데이터 업데이트 시\n재학습 필요", "중간\n(학습 데이터 부족\n가능성)"],
            ["Long Context", "전체 API 문서를\n긴 컨텍스트에 삽입", "구현 가장 단순\n파이프라인 불필요", "토큰 비용 높음\n모델 제한\n(Claude 1M 토큰 등)", "Phase 2 초기\nPoC에 적합"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] RAG vs Fine-Tuning vs Long Context 2026 — dev.to/pockit_tools")
    add_ref(doc, "[Ref] Vibe Coding 가이드 2026 — sitepoint.com/vibe-coding-complete-developers-guide")

    doc.add_heading("5.4  XM_API 특화 코드 검증 파이프라인 (방향성)", level=2)
    add_table(doc,
        ["단계", "검증 내용", "실패 시"],
        [
            ["1. 구조 검증", "User_Setup(), User_Loop() 존재 확인", "LLM 재생성 요청"],
            ["2. API 화이트리스트", "XM_API 함수만 호출 허용\n(HAL, RTOS 직접 호출 거부)", "위반 API 식별 + 대체 API 제안"],
            ["3. 금지 패턴 검출", "malloc, free, __disable_irq,\ntaskENTER_CRITICAL 등", "금지 패턴 제거 + 안전 대안 제시"],
            ["4. 컴파일 테스트", "클라우드 빌드 실행\n(GCC 에러/경고)", "에러 메시지로 LLM 재생성"],
        ], first_col_bold=True
    )

    doc.add_heading("5.5  다음 단계", level=2)
    add_bullet(doc, "Phase 1 개발 기간 중 방법론 선정 완료 (A/B/C/D 중 결정)")
    add_bullet(doc, "GPU 서버 가용성 확인 (A5000 4장 활용 가능성)")
    add_bullet(doc, "외부 API 비용 추정 (요청당 평균 토큰, 월간 사용량)")
    add_bullet(doc, "Phase 2 착수 시점에 별도 AI 기술사양서 작성 (본 문서와 동일 양식)")

    # ================================================================
    # Section 6
    # ================================================================
    doc.add_heading("6. 개발 우선순위 및 일정", level=1)
    add_table(doc,
        ["Phase", "내용", "예상 기간", "상태"],
        [
            ["Phase 0", "인프라\n(libXM_Full.a + Docker SDK + XM_API 스키마)", "4~6주", "기획 확정"],
            ["Phase 0 (병렬)", "AGR_BOOT V2 Bootloader", "4~6주", "별도 기획서 완료"],
            ["Phase 1", "Sandbox + 에디터 + 클라우드 빌드\n+ FW 업로드 연동", "8~10주", "기획 확정"],
            ["Phase 2", "AI LLM 코드 생성", "미정", "기획 심화 중"],
        ], first_col_bold=True
    )
    add_body(doc,
        "Phase 0과 Phase 1은 순차 진행하되, Bootloader는 Phase 0과 병렬 진행 가능.",
        bold=True
    )

    # ================================================================
    # Section 7
    # ================================================================
    doc.add_heading("7. 담당 분야 (업무 배분 기준)", level=1)
    add_table(doc,
        ["분야", "핵심 기술", "산출물"],
        [
            ["Static Library 빌드", "CMake, ARM GCC,\n링커 스크립트", "libXM_Full.a\n+ Docker SDK 패키지"],
            ["클라우드 컴파일 서비스", "FastAPI, Docker,\nARM GCC", "REST API\n+ Docker 이미지"],
            ["웹 에디터 (Frontend)", "Monaco Editor,\nTypeScript, React", "코드 에디터\n+ 파일 탐색기"],
            ["Sandbox", "React, JSON,\n예제 관리", "예제 갤러리 UI\n+ 메타데이터 JSON"],
            ["FW 업로드 연동", "Web Serial,\nTypeScript", "FTP Host Engine\n(Spec v1.1 기반)"],
            ["Bootloader", "C, STM32 HAL,\nFlash", "AGR_BOOT V2\n(별도 기획서)"],
            ["XM_API 스키마", "JSON, XM_API 분석", "API 정의 JSON"],
        ], first_col_bold=True
    )

    # ================================================================
    # Section 8
    # ================================================================
    doc.add_heading("8. 리스크 및 제약 조건", level=1)
    add_table(doc,
        ["리스크", "영향", "대응"],
        [
            ["Web Serial API\n브라우저 제한", "Chrome/Edge만 지원\n(Safari, Firefox 미지원)", "지원 브라우저 명시\n비지원 시 경고 UI 표시\nFeature Detection 적용"],
            ["클라우드 빌드 보안", "사용자 코드 인젝션\n리소스 남용", "Docker 격리 (네트워크 차단)\n빌드 타임아웃 30초\n파일 크기 제한 500KB"],
            ["libXM_Full.a\n버전 관리", "XM FW 업데이트 시\n.a 재빌드 필요", ".a 빌드 CI/CD 자동화\nDocker 이미지 태그 버전 관리"],
            ["Toolchain 호환성", ".a 빌드와 사용자 코드\n컴파일 버전 불일치", "Docker 이미지에\narm-none-eabi-gcc 고정 버전"],
            ["HTTPS 요건", "Web Serial API는\nHTTPS 필수", "개발: localhost 허용\n배포: SSL 인증서 설정"],
            ["AI LLM 방법론\n미확정", "Phase 2 착수 시점\n불투명", "Phase 1 기간 중\n비교 검토 완료\nGPU 가용성 확인"],
            ["Hardware Hallucination\n(AI 생성 코드)", "존재하지 않는 API 호출\n잘못된 HW 접근", "API 화이트리스트 필수\n4단계 검증 파이프라인"],
        ], first_col_bold=True
    )

    # ================================================================
    # Section 9
    # ================================================================
    doc.add_heading("9. 레퍼런스", level=1)

    doc.add_heading("9.1  내부 문서", level=2)
    add_table(doc,
        ["문서", "위치"],
        [
            ["PhAI Studio FW Upload 기술사양서 v1.1", "docs/bootloader/"],
            ["XM V2 Bootloader Plan", "plan 파일"],
            ["Total Data Packet 설계사양서", "docs/total_data_packet/"],
            ["XM10 SDK 릴리즈 레포", "Extension_Module/"],
        ], first_col_bold=True
    )

    doc.add_heading("9.2  외부 레퍼런스", level=2)
    add_table(doc,
        ["영역", "레퍼런스", "URL"],
        [
            ["클라우드 빌드", "ESP-IDF Docker Image", "docs.espressif.com/projects/esp-idf"],
            ["클라우드 빌드", "Arduino Cloud Builder\n(빌드 50% 단축)", "blog.arduino.cc/2024/12/13"],
            ["Static Library", "Mbed OS .a 호환성", "os.mbed.com/questions/83007"],
            ["Monaco Editor", "Monaco 통합 가이드", "blog.spectralcore.com/\nintegrating-monaco-editor"],
            ["Monaco + LSP", "Monaco + LSP 설정", "dev.to/__4f1641"],
            ["Web Serial API", "MDN 공식 문서", "developer.mozilla.org/en-US/\ndocs/Web/API/Web_Serial_API"],
            ["Web Serial FW", "Bleskomat Web Serial", "github.com/bleskomat/\nbleskomat-web-serial"],
            ["Web Serial FW", "O.MG WebFlasher", "github.com/O-MG/WebFlasher"],
            ["AI 임베디드", "Vibe Coding 가이드 2026", "sitepoint.com/vibe-coding"],
            ["AI 임베디드", "RAG vs Fine-Tuning 2026", "dev.to/pockit_tools"],
            ["AI 제약", "AI in Embedded Systems", "wedolow.com/resources/\nvibe-coding-embedded"],
            ["Sandbox", "CodeSandbox 템플릿", "codesandbox.io/docs/sdk/templates"],
            ["플랫폼", "Keil Studio Cloud", "keil.arm.com"],
            ["플랫폼", "PlatformIO Remote", "docs.platformio.org/en/latest/\nplus/pio-remote.html"],
            ["플랫폼", "Wokwi STM32 Simulator", "wokwi.com/stm32"],
            ["플랫폼", "Renesas 365 Cloud", "electronicsmedia.info/2026/03/11"],
        ], first_col_bold=True
    )

    doc.add_heading("9.3  기술 판단 (확정)", level=2)
    add_table(doc,
        ["결정", "내용", "근거"],
        [
            ["Static Library", "단일 libXM_Full.a\n--whole-archive + --gc-sections", "Arduino Cloud 동일 전략\n빌드 3~8초 목표"],
            ["FW 업로드 프로토콜", "AGR FTP over USB CDC", "FW Upload Spec v1.1\nDFU 대비 장점 참조"],
            ["FW 전송 경로", "서버 미경유\nBrowser → USB 직접", "Bleskomat, O.MG 검증\n오프라인 동작 보장"],
            ["빌드 산출물", "packaged.bin\n(AGR_FwInfo_t 포함)", "BL FTP가 이 형식 기대\n호환성 사전 검증 활용"],
            ["Backend 역할", "메타데이터 Optional Add-On\n(FW 전송과 무관)", "FW Upload Spec v1.1\nSection 8 원칙"],
            ["빌드 서버", "기존 SLURM 서버에\nDocker 컨테이너", "ESP-IDF Docker 패턴\n추가 인프라 불필요"],
            ["개발 우선순위", "Sandbox + 직접 코딩 먼저\nAI LLM은 기획 후 착수", "Hardware Hallucination\n리스크 관리"],
        ], first_col_bold=True
    )

    add_ref(doc, "[Ref] PhAI Studio FW Upload 기술사양서 v1.1 — docs/bootloader/")
    add_ref(doc, "[Ref] XM V2 Bootloader Plan — plan 파일")


# ============================================================
# English Version
# ============================================================
def build_xm_studio_spec_en(doc):
    """XM-Studio Web IDE Technical Specification (English)."""

    add_title_page(doc,
        "XM-Studio Web IDE\nTechnical Specification",
        "XM10 Embedded System Web-Based Development Environment — Tech Spec & Development Plan",
        "v1.0  ·  2026-03-11  ·  Angel Robotics"
    )

    add_info_cards(doc, [
        ("Target", "PhAI Studio Web + XM10 FW", BG_INFO_BLUE),
        ("Core Principle", "API-Constrained IDE (XM_API)", BG_INFO_GRAY),
        ("Build", "Docker + libXM_Full.a", BG_INFO_GREEN),
        ("FW Transfer", "Web Serial FTP (No Server Relay)", BG_INFO_GREEN),
    ])
    add_info_cards(doc, [
        ("Priority", "Sandbox + Direct Coding First", BG_INFO_BLUE),
        ("AI LLM", "Planning Stage (Not Implemented)", BG_INFO_AMBER),
        ("Version", "v1.0 (Technical Specification)", BG_INFO_GRAY),
        ("Status", "Plan Finalized, Dev Prep", BG_INFO_GRAY),
    ])

    toc_items = [
        "1. Overview",
        "2. System Architecture",
        "3. Phase 0: Infrastructure Spec",
        "4. Phase 1: Core Features (Sandbox + Direct Coding)",
        "5. Phase 2: AI LLM Code Generation (Planning Stage)",
        "6. Development Priority & Schedule",
        "7. Role Allocation (Task Distribution)",
        "8. Risks & Constraints",
        "9. References",
    ]
    _add_toc(doc, toc_items, "en")

    # === Section 1 ===
    doc.add_heading("1. Overview", level=1)

    doc.add_heading("1.1  Purpose", level=2)
    add_body(doc,
        "This document defines the technical specification and development plan for XM-Studio, "
        "a web-based development environment for the XM10 embedded system. "
        "It enables users to write, compile, and upload XM10 user code using only a web browser, "
        "without requiring STM32CubeIDE or physical debugging tools."
    )
    add_body(doc,
        "Adopting the Roblox Studio paradigm: users focus solely on business logic "
        "(control algorithms, data collection, etc.) while the platform provides all infrastructure "
        "(HAL, RTOS, drivers)."
    )

    doc.add_heading("1.2  Scope", level=2)
    add_bullet(doc, "PhAI Studio Frontend (Next.js + TypeScript + Monaco Editor)")
    add_bullet(doc, "Cloud Compile Service (FastAPI + Docker + ARM GCC)")
    add_bullet(doc, "XM10 SDK-based Static Library (libXM_Full.a)")
    add_bullet(doc, "FW Upload (Web Serial FTP — aligned with FW Upload Spec v1.1)")
    add_bullet(doc, "AI LLM Code Generation (planning stage, Phase 2)")

    doc.add_heading("1.3  Core Design Principles", level=2)
    add_table(doc,
        ["Principle", "Description"],
        [
            ["XM_API Facade Constraint", "Users can only call XM_API functions. Direct HAL/RTOS/driver calls prohibited"],
            ["Direct Transfer (No Server)", "FW binary sent directly Browser -> USB CDC (FW Upload Spec v1.1)"],
            ["Static Library Build", "Single libXM_Full.a + compile only user code for minimal build time"],
            ["Dev Priority", "Sandbox + direct coding first, AI LLM after server infrastructure"],
            ["packaged.bin Output", "Cloud build generates packaged.bin including AGR_FwInfo_t header"],
        ], first_col_bold=True
    )

    doc.add_heading("1.4  Reference Documents", level=2)
    add_table(doc,
        ["Document", "Location", "Description"],
        [
            ["PhAI Studio FW Upload\nSpec v1.1", "docs/bootloader/", "FTP protocol, AGR_FwInfo_t, upload workflow"],
            ["XM V2 Bootloader Plan", "plan file", "BL design (Flash Layout, Core, Transport)"],
            ["Total Data Packet Spec", "docs/total_data_packet/", "Module ID scheme, USB CDC streaming"],
            ["XM10 SDK Release Repo", "Extension_Module/", ".a file, XM_API, .cproject"],
        ], first_col_bold=True
    )

    doc.add_heading("1.5  Glossary", level=2)
    add_table(doc,
        ["Term", "Definition"],
        [
            ["XM_API", "XM10 user facade API (~80 functions). Called from User_Setup/User_Loop"],
            ["libXM_Full.a", "Single static library: Core + HAL + CMSIS + FreeRTOS + Middleware + XM_FW"],
            ["packaged.bin", "AGR_FwInfo_t header (1KB) + App Binary (max 767KB). Format expected by bootloader"],
            ["AGR FTP", "Angel Robotics Firmware Transfer Protocol over USB CDC (PhAI V2.2 packets)"],
            ["Sandbox", "Pre-configured example code gallery for instant load/edit/build"],
            ["Vibe Coding", "Development paradigm where LLM generates code from natural language prompts"],
        ], first_col_bold=True
    )

    # === Section 2 ===
    doc.add_heading("2. System Architecture", level=1)

    doc.add_heading("2.1  Overall Architecture", level=2)
    add_body(doc,
        "XM-Studio consists of a 3-stage pipeline: code editing (browser) -> cloud compilation (server) -> "
        "FW upload (browser -> USB direct transfer). Compilation goes through the server, "
        "but FW transfer bypasses the server entirely."
    )
    add_code_block(doc,
        "=== XM-Studio Architecture ===\n"
        "\n"
        "  [User Browser]\n"
        "       | Monaco Editor (C code editing)\n"
        "       | Sandbox (load/edit examples)\n"
        "       v\n"
        "  [PhAI Studio Frontend (Next.js)]\n"
        "       | POST /api/compile (multi-file source)\n"
        "       v\n"
        "  [Cloud Compile Server (FastAPI + Docker)]\n"
        "       | Docker Container: arm-none-eabi-gcc\n"
        "       | CMake + libXM_Full.a linking\n"
        "       | objcopy -> raw .bin -> agr_fw_packager -> packaged.bin\n"
        "       v\n"
        "  [packaged.bin download -> browser memory]\n"
        "       | FTP Host Engine (TypeScript)\n"
        "       v\n"
        "  [Web Serial API] ---- USB CDC, 921600 baud\n"
        "       |\n"
        "  ======== USB Cable (Physical) ========\n"
        "       |\n"
        "  [XM10 Bootloader -> Flash Write]\n"
        "\n"
        "  Key: Compilation via server / FW transfer NOT via server (direct USB)"
    )

    doc.add_heading("2.2  Reference Platform Comparison", level=2)
    add_table(doc,
        ["Platform", "Editor", "Build", "Upload", "AI", "Differentiation"],
        [
            ["Arduino Web Editor", "Monaco", "Cloud\n(Pre-compiled)", "Web Serial\n(esptool)", "Copilot\n(partial)", "No API constraint\n(full HAL exposed)"],
            ["Keil Studio Cloud", "VS Code Web", "Cloud\n(ARMC6)", "CMSIS-DAP", "None", "ARM only\nbrowser limits"],
            ["PlatformIO Remote", "VS Code", "Remote Agent", "Via Agent", "None", "Agent install\nrequired"],
            ["Wokwi", "Monaco", "Local/CI", "Simulator", "None", "Simulator only\n(no real upload)"],
            ["XM-Studio (Ours)", "Monaco", "Docker +\nlibXM_Full.a", "Web Serial FTP\n(packaged.bin)", "RAG\n(Phase 2)", "API-constrained IDE\npackaged.bin FTP\nBackup-First safety"],
        ], first_col_bold=True
    )

    # === Section 3 ===
    doc.add_heading("3. Phase 0: Infrastructure Spec", level=1)

    doc.add_heading("3.1  libXM_Full.a — Single Static Library", level=2)
    add_body(doc,
        "All non-user code in XM10 FW is consolidated into a single .a file. "
        "The cloud build only compiles user code and links against this library. "
        "Same strategy as Arduino Cloud's pre-compiled core (50% build time reduction)."
    )
    add_table(doc,
        ["Layer", "Contents", "Est. .o Count"],
        [
            ["Core", "core_process.c, system_init.c, etc.", "~10"],
            ["HAL", "STM32H7 HAL drivers", "~40"],
            ["CMSIS", "system_stm32h7xx.c, DSP library", "~15"],
            ["FreeRTOS", "RTOS kernel + CMSIS-OS2 wrapper", "~20"],
            ["Middleware", "FatFs, USB Host/Device", "~30"],
            ["Compatible", "stm32h7xx_compatible.c, etc.", "~5"],
            ["XM_FW", "XM_API impl, IOIF, Devices, Services", "~100"],
        ], first_col_bold=True
    )
    add_body(doc, "Compile & Link Strategy:", bold=True)
    add_table(doc,
        ["Item", "Setting", "Reason"],
        [
            ["Compile Flags", "-ffunction-sections\n-fdata-sections", "Per-function/data sections\nfor dead code elimination"],
            ["Linker: --whole-archive", "Applied to libXM_Full.a", "Force include all .o so extern\nUser_Setup/User_Loop resolve"],
            ["Linker: --gc-sections", "Final link", "Remove unused functions from\n--whole-archive included .o files"],
            ["Toolchain Fixed", "arm-none-eabi-gcc\n12.3.rel1", "Same version for .a build and\nuser code (Mbed compatibility)"],
        ], first_col_bold=True
    )

    doc.add_heading("3.2  Docker SDK Package", level=2)
    add_code_block(doc,
        "Docker SDK Package Layout:\n"
        "\n"
        "  /opt/xm-sdk/\n"
        "  +-- lib/libXM_Full.a\n"
        "  +-- include/XM_API/          # xm_api.h + sub-headers\n"
        "  +-- include/internal/        # Transitive dependency headers (~10)\n"
        "  +-- include/rtos/            # cmsis_os2.h\n"
        "  +-- linker/STM32H743XIHx_FLASH.ld\n"
        "  +-- startup/startup_stm32h743xihx.s\n"
        "  +-- cmake/cloud_build.cmake\n"
        "  +-- tools/agr_fw_packager.py # raw .bin -> packaged.bin"
    )

    doc.add_heading("3.3  XM_API JSON Schema", level=2)
    add_body(doc,
        "XM_API function signatures, parameters, return values, and type definitions "
        "structured as JSON for Monaco IntelliSense and future AI RAG data."
    )
    add_table(doc,
        ["Item", "Content"],
        [
            ["Target Functions", "~80 XM_API public functions"],
            ["Included Types", "enum, struct, typedef (XM_API exposed only)"],
            ["Usage", "Monaco registerCompletionItemProvider (Phase 1)\nAI RAG data (Phase 2)"],
            ["Generation", "XM_API header parsing script (Python)"],
        ], first_col_bold=True
    )

    doc.add_heading("3.4  packaged.bin Build Pipeline", level=2)
    add_code_block(doc,
        "Build Pipeline:\n"
        "\n"
        "  [User Code (*.c, *.h)]\n"
        "       | arm-none-eabi-gcc -c\n"
        "       v\n"
        "  [User .o files]\n"
        "       | Linker: --whole-archive libXM_Full.a + --gc-sections\n"
        "       v\n"
        "  [firmware.elf] -> objcopy -> [firmware.bin] (raw)\n"
        "       | agr_fw_packager.py\n"
        "       v\n"
        "  [packaged.bin] = AGR_FwInfo_t(1KB) + firmware.bin(max 767KB)\n"
        "       -> Browser download -> FTP Host Engine -> Web Serial -> XM BL"
    )

    # === Section 4 ===
    doc.add_heading("4. Phase 1: Core Features (Sandbox + Direct Coding)", level=1)
    add_body(doc,
        "Phase 1 = Sandbox + Direct Typing + Cloud Build + FW Upload. AI LLM separated into Phase 2.",
        bold=True
    )

    doc.add_heading("4.1  Monaco Editor — C Language IDE", level=2)
    add_table(doc,
        ["Feature", "Implementation", "Notes"],
        [
            ["Syntax Highlighting", "Monaco built-in C support", "Native"],
            ["XM_API Autocomplete", "registerCompletionItemProvider\n+ JSON schema", "Custom (Monaco lacks\nbuilt-in C LSP)"],
            ["Compile Error Markers", "GCC output parsing\n-> setModelMarkers", "Server-dependent"],
            ["Multi-file/Folder", "File explorer (tree view)\nmulti-tab, new file/folder", "Custom UI"],
            ["User_Setup/Loop Check", "Pre-build client validation\n(regex match)", "Prevent build failure"],
        ], first_col_bold=True
    )
    add_body(doc, "Monaco C Language Limitations:", bold=True)
    add_bullet(doc, "Monaco provides only syntax highlighting for C natively — no real-time error validation")
    add_bullet(doc, "Error detection depends on cloud build GCC output (Save -> Build -> Show errors cycle)")
    add_bullet(doc, "Future: Clangd LSP via WebSocket for real-time error display (Phase 2+ consideration)")

    doc.add_heading("4.2  Sandbox (Example Gallery)", level=2)
    add_table(doc,
        ["Item", "Content"],
        [
            ["Source", "52 existing examples (ARC_ExtensionBoard/.../Examples/)"],
            ["Categories", "Beginner(6), Elementary(12), Intermediate(15), Advanced(19)"],
            ["Metadata", "JSON: id, title, category, difficulty, description, files[], required_hw"],
            ["Workflow", "Select example -> Load in Editor -> Modify -> Build -> Upload"],
            ["Per-example README", "API list used, behavior description, HW wiring guide"],
        ], first_col_bold=True
    )

    doc.add_heading("4.3  Cloud Compile Service", level=2)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ["POST", "/api/compile", "Submit multi-file source -> return packaged.bin"],
            ["GET", "/api/compile/{job_id}/status", "Build status (pending/building/done/error)"],
            ["GET", "/api/compile/{job_id}/binary", "Download packaged.bin"],
            ["GET", "/api/compile/{job_id}/errors", "GCC errors/warnings (JSON)"],
        ], first_col_bold=True
    )
    add_body(doc, "Build Pipeline:", bold=True)
    add_table(doc,
        ["Step", "Action", "Est. Time"],
        [
            ["1. Receive", "JSON -> temp directory files", "< 0.1s"],
            ["2. CMake", "Auto-configure via cloud_build.cmake", "< 0.5s"],
            ["3. Compile", "gcc -c (user code only)", "1-3s"],
            ["4. Link", "libXM_Full.a + user .o + startup.o", "1-2s"],
            ["5. objcopy", ".elf -> raw .bin", "< 0.1s"],
            ["6. Package", "agr_fw_packager.py -> packaged.bin", "< 0.1s"],
            ["7. Response", "packaged.bin + memory usage report", "< 0.1s"],
        ], first_col_bold=True
    )
    add_body(doc, "Target build time: 3~8 seconds", bold=True)

    doc.add_heading("4.4  FW Upload Integration (Web Serial FTP)", level=2)
    add_body(doc,
        "Detailed FW upload protocol specification is in a separate document "
        "(PhAI Studio FW Upload Spec v1.1). This section covers XM-Studio integration only.",
        bold=True
    )
    add_table(doc,
        ["Item", "Content"],
        [
            ["Build Output", "packaged.bin (cloud build includes AGR_FwInfo_t header)"],
            ["Transfer Path", "Browser memory -> FTP Host Engine -> Web Serial -> USB CDC -> XM BL"],
            ["No Server Relay", "packaged.bin transferred directly via USB (FW Upload Spec v1.1)"],
            ["Compatibility", "AGR_FwInfo_t header: target_module, hw_rev_min/max pre-validation"],
            ["Upload 5 Stages", "Backup -> Erase -> Transfer -> Verify -> Commit"],
            ["HTTPS Required", "Web Serial API requires secure context (HTTPS or localhost)"],
        ], first_col_bold=True
    )

    # === Section 5 ===
    doc.add_heading("5. Phase 2: AI LLM Code Generation (Planning Stage)", level=1)
    add_body(doc,
        "Planning phase during Phase 1 development and server infrastructure setup. "
        "Implementation starts after Phase 1 completion with a separate AI tech spec document.",
        bold=True
    )

    doc.add_heading("5.1  Technical Approach Comparison", level=2)
    add_table(doc,
        ["Approach", "Architecture", "Pros", "Cons", "Cost"],
        [
            ["A: External API\n(Claude/GPT/Gemini)", "User -> Backend\n-> External API", "Latest models\nHigh performance\nFast deployment", "Token-based cost\nExternal dependency\nNetwork required", "Pay-per-use"],
            ["B: Self-hosted\nOpen Source", "User -> Backend\n-> Local GPU", "Fixed cost\nNo data leakage\nOffline capable", "GPU server needed\nModel limitations\nMaintenance", "Fixed\n(GPU server)"],
            ["C: Hybrid", "Simple = Local\nComplex = External", "Cost/perf balance", "Implementation complexity\nRouting logic needed", "Mixed"],
            ["D: User API Key", "User provides\nown API key", "Zero Angel cost\nUser choice", "High UX barrier\nComplex setup", "User-paid"],
        ], first_col_bold=True
    )

    doc.add_heading("5.2  Hardware Hallucination Risk", level=2)
    add_body(doc,
        "The core risk of embedded AI code generation is 'Hardware Hallucination' — "
        "LLMs generating non-existent APIs, wrong pin numbers, or inappropriate "
        "peripheral initialization sequences."
    )
    add_table(doc,
        ["Risk", "Description", "XM-Studio Mitigation"],
        [
            ["Non-existent API calls", "LLM generates HAL_GPIO_WritePin\n(not in XM_API)", "API whitelist validation\n(only 80 XM_API functions)"],
            ["Invalid parameters", "Out-of-range values,\ninvalid pin numbers", "Parameter range check\n+ compile test"],
            ["Timing unawareness", "Blocking calls in\nreal-time control loop", "Forbidden pattern detection\n(malloc, __disable_irq)"],
            ["Structure violation", "Code outside\nUser_Setup/User_Loop", "Structure validation\n(required function check)"],
        ], first_col_bold=True
    )

    doc.add_heading("5.3  RAG vs Fine-Tuning vs Long Context", level=2)
    add_table(doc,
        ["Method", "Principle", "Pros", "Cons", "XM-Studio Fit"],
        [
            ["RAG", "Retrieve relevant data\nat query time", "No model change\nEasy data updates\nCost effective", "Search pipeline needed\nDepends on retrieval", "High\n(80 APIs + 52 examples\n= good size)"],
            ["Fine-Tuning", "Adjust model weights\nwith domain data", "Domain-specialized\nNo inference overhead", "High training cost\nRetrain on updates", "Medium\n(limited training\ndata)"],
            ["Long Context", "Insert full API docs\nin long context", "Simplest implementation\nNo pipeline needed", "High token cost\nModel limits", "Good for Phase 2\ninitial PoC"],
        ], first_col_bold=True
    )

    # === Section 6 ===
    doc.add_heading("6. Development Priority & Schedule", level=1)
    add_table(doc,
        ["Phase", "Content", "Est. Duration", "Status"],
        [
            ["Phase 0", "Infrastructure\n(libXM_Full.a + Docker SDK + JSON Schema)", "4~6 weeks", "Plan Finalized"],
            ["Phase 0 (Parallel)", "AGR_BOOT V2 Bootloader", "4~6 weeks", "Separate Spec Done"],
            ["Phase 1", "Sandbox + Editor + Cloud Build\n+ FW Upload Integration", "8~10 weeks", "Plan Finalized"],
            ["Phase 2", "AI LLM Code Generation", "TBD", "Planning In Progress"],
        ], first_col_bold=True
    )

    # === Section 7 ===
    doc.add_heading("7. Role Allocation (Task Distribution)", level=1)
    add_table(doc,
        ["Area", "Key Technologies", "Deliverables"],
        [
            ["Static Library Build", "CMake, ARM GCC,\nLinker Script", "libXM_Full.a\n+ Docker SDK Package"],
            ["Cloud Compile Service", "FastAPI, Docker,\nARM GCC", "REST API\n+ Docker Image"],
            ["Web Editor (Frontend)", "Monaco Editor,\nTypeScript, React", "Code Editor\n+ File Explorer"],
            ["Sandbox", "React, JSON,\nExample Management", "Example Gallery UI\n+ Metadata JSON"],
            ["FW Upload Integration", "Web Serial,\nTypeScript", "FTP Host Engine\n(Spec v1.1 based)"],
            ["Bootloader", "C, STM32 HAL,\nFlash", "AGR_BOOT V2\n(Separate Spec)"],
            ["XM_API Schema", "JSON, XM_API Analysis", "API Definition JSON"],
        ], first_col_bold=True
    )

    # === Section 8 ===
    doc.add_heading("8. Risks & Constraints", level=1)
    add_table(doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["Web Serial API\nBrowser Limits", "Chrome/Edge only\n(Safari, Firefox N/A)", "Document supported browsers\nWarning UI, Feature Detection"],
            ["Cloud Build Security", "Code injection\nResource abuse", "Docker isolation (no network)\nBuild timeout 30s\nFile size limit 500KB"],
            ["libXM_Full.a\nVersioning", "Rebuild needed on\nXM FW update", "CI/CD automation\nDocker image tag versioning"],
            ["Toolchain Compat", ".a build vs user code\nversion mismatch", "Fixed arm-none-eabi-gcc\nin Docker image"],
            ["HTTPS Requirement", "Web Serial API\nneeds HTTPS", "Dev: localhost OK\nProd: SSL certificate"],
            ["AI LLM Methodology\nUndecided", "Phase 2 start\nuncertain", "Complete comparison\nduring Phase 1\nCheck GPU availability"],
            ["Hardware Hallucination\n(AI generated code)", "Non-existent API calls\nWrong HW access", "API whitelist mandatory\n4-stage verification pipeline"],
        ], first_col_bold=True
    )

    # === Section 9 ===
    doc.add_heading("9. References", level=1)

    doc.add_heading("9.1  Internal Documents", level=2)
    add_table(doc,
        ["Document", "Location"],
        [
            ["PhAI Studio FW Upload Spec v1.1", "docs/bootloader/"],
            ["XM V2 Bootloader Plan", "plan file"],
            ["Total Data Packet Spec", "docs/total_data_packet/"],
            ["XM10 SDK Release Repo", "Extension_Module/"],
        ], first_col_bold=True
    )

    doc.add_heading("9.2  External References", level=2)
    add_table(doc,
        ["Area", "Reference", "URL"],
        [
            ["Cloud Build", "ESP-IDF Docker Image", "docs.espressif.com/projects/esp-idf"],
            ["Cloud Build", "Arduino Cloud Builder\n(50% faster builds)", "blog.arduino.cc/2024/12/13"],
            ["Static Library", "Mbed OS .a compat", "os.mbed.com/questions/83007"],
            ["Monaco Editor", "Integration Guide", "blog.spectralcore.com/\nintegrating-monaco-editor"],
            ["Web Serial API", "MDN Official Docs", "developer.mozilla.org/en-US/\ndocs/Web/API/Web_Serial_API"],
            ["Web Serial FW", "Bleskomat Web Serial", "github.com/bleskomat"],
            ["Web Serial FW", "O.MG WebFlasher", "github.com/O-MG/WebFlasher"],
            ["AI Embedded", "Vibe Coding Guide 2026", "sitepoint.com/vibe-coding"],
            ["AI Embedded", "RAG vs Fine-Tuning 2026", "dev.to/pockit_tools"],
            ["AI Constraints", "AI in Embedded Systems", "wedolow.com/resources/vibe-coding"],
            ["Sandbox", "CodeSandbox Templates", "codesandbox.io/docs/sdk/templates"],
            ["Platform", "Keil Studio Cloud", "keil.arm.com"],
            ["Platform", "PlatformIO Remote", "docs.platformio.org"],
            ["Platform", "Wokwi STM32", "wokwi.com/stm32"],
            ["Platform", "Renesas 365 Cloud", "electronicsmedia.info/2026/03/11"],
        ], first_col_bold=True
    )

    doc.add_heading("9.3  Technical Decisions (Finalized)", level=2)
    add_table(doc,
        ["Decision", "Content", "Rationale"],
        [
            ["Static Library", "Single libXM_Full.a\n--whole-archive + --gc-sections", "Same as Arduino Cloud\nTarget: 3~8s build"],
            ["FW Upload Protocol", "AGR FTP over USB CDC", "FW Upload Spec v1.1\nAdvantages over DFU"],
            ["FW Transfer Path", "No server relay\nBrowser -> USB direct", "Bleskomat, O.MG proven\nOffline operation"],
            ["Build Output", "packaged.bin\n(with AGR_FwInfo_t)", "BL FTP expects this format\nPre-validation support"],
            ["Backend Role", "Metadata optional add-on\n(not in FW path)", "FW Upload Spec v1.1\nSection 8 principle"],
            ["Build Server", "Docker on existing\nSLURM server", "ESP-IDF Docker pattern\nNo additional infra"],
            ["Dev Priority", "Sandbox + Coding first\nAI LLM after planning", "Hardware Hallucination\nrisk management"],
        ], first_col_bold=True
    )

    add_ref(doc, "[Ref] PhAI Studio FW Upload Spec v1.1 — docs/bootloader/")
    add_ref(doc, "[Ref] XM V2 Bootloader Plan — plan file")


def main():
    doc_kr = setup_document(FONT_KR)
    build_xm_studio_spec_kr(doc_kr)
    kr_path = os.path.join(OUTPUT_DIR, "XM_Studio_Web_IDE_기술기획서_v1.0.docx")
    doc_kr.save(kr_path)
    print(f"[OK] Korean: {kr_path}")

    doc_en = setup_document(FONT_KR)
    build_xm_studio_spec_en(doc_en)
    en_path = os.path.join(OUTPUT_DIR, "XM_Studio_Web_IDE_Spec_v1.0.docx")
    doc_en.save(en_path)
    print(f"[OK] English: {en_path}")


if __name__ == "__main__":
    main()
