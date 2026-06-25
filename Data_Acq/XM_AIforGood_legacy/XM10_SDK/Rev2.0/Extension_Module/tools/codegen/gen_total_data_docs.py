"""
Total Data Packet 설계문서 Word(.docx) 생성기.

PhAI_Studio_기술사양서_v1.1의 Section 3 "Host Subscribe Model"의 연속 문서로,
XM10 Total Data Packet 구현 사양을 상세히 기술합니다.

Usage:
    python docs/total_data_packet/gen_total_data_docs.py

Dependencies:
    pip install python-docx pyyaml
"""
import sys
import os

# gen_docx_v2_base.py 경로 추가
BASE_DIR = r"C:\AGR_HyundoKim\02_PhAI\PhAI_Web\Docs\Plan\PhAI_Studio_기획"
sys.path.insert(0, BASE_DIR)
sys.stdout.reconfigure(encoding="utf-8")

from gen_docx_v2_base import (
    setup_document, add_title_page, add_info_cards, add_table, add_code_block,
    add_body, add_bullet, add_ref, FONT_KR,
    BG_INFO_BLUE, BG_INFO_GRAY, BG_INFO_GREEN, BG_INFO_AMBER,
)

F = FONT_KR
OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def build_total_data_packet_doc(doc):
    """00_Total_Data_Packet_Design — 전체 아키텍처 설계."""

    # === Title Page ===
    add_title_page(doc,
        "XM10 Total Data Packet\n설계 사양서",
        "PhAI Studio 기술사양서 v1.1 — Section 3 Host Subscribe Model 구현 사양",
        "v1.0  ·  2026-03-10  ·  Angel Robotics"
    )

    add_info_cards(doc, [
        ("기준 문서", "기술사양서 v1.1 Section 3", BG_INFO_BLUE),
        ("적용 대상", "XM10 Rev 1.1 FW", BG_INFO_GRAY),
        ("패킷 크기", "~425 Bytes @ 1kHz", BG_INFO_GREEN),
        ("USB FS 점유", "~55.4% (안전 마진 내)", BG_INFO_GREEN),
    ])
    add_info_cards(doc, [
        ("XM FW", "System 자동 전송 (Module ID 0x20)", BG_INFO_BLUE),
        ("User Custom", "Module ID 0xF0~0xFE (선택적)", BG_INFO_AMBER),
        ("Data Map", "YAML SSOT → C/TS Code-Gen", BG_INFO_GRAY),
        ("phai-studio", "Phase 2 (별도 구현 예정)", BG_INFO_GRAY),
    ])

    # === TOC ===
    doc.add_heading("목 차", level=1)
    toc_items = [
        "1. 개요 및 배경",
        "2. 설계 결정 요약",
        "3. Module ID 체계",
        "4. Total Data Packet 구조",
        "5. 대역폭 분석",
        "6. XM FW 데이터 흐름",
        "7. Data Map 파이프라인 (YAML → Code-Gen)",
        "8. User Custom Data API",
        "9. phai-studio 구현 요구사항",
        "10. API 변경 요약",
        "11. 리스크 및 제약 조건",
    ]
    from docx.shared import Pt
    from gen_docx_v2_base import CLR_GRAY
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("   " + item)
        r.font.name = F
        r.font.size = Pt(10)
        r.font.color.rgb = CLR_GRAY
    doc.add_page_break()

    # ========== Section 1 ==========
    doc.add_heading("1. 개요 및 배경", level=1)

    doc.add_heading("1.1  기술사양서 Section 3과의 관계", level=2)
    add_body(doc,
        "PhAI Studio 기술사양서 v1.1의 Section 3 \"데이터 교환 모델 — Host Subscribe Model\"에서 "
        "Full Broadcast + Host-Side Filtering 방식을 채택하였다. "
        "본 문서는 해당 모델의 XM10 FW 측 구체적 구현 사양을 정의한다."
    )
    add_body(doc,
        "기존 기술사양서에서 \"COMBINED 패킷(Module ID 0x10)에 필드를 추가\"하는 방식이었으나, "
        "이를 발전시켜 전용 Module ID(0x20)의 Total Data Packet으로 재설계하였다. "
        "이를 통해 사용자 코드 없이 System이 모든 센서 raw 데이터를 자동 전송하며, "
        "사용자는 phai-studio에서 원하는 채널만 선택하여 모니터링/저장할 수 있다."
    )

    doc.add_heading("1.2  핵심 설계 원칙", level=2)
    add_bullet(doc, "Stream-All, Filter-on-Host: XM은 전부 보냄, PC에서 선택")
    add_bullet(doc, "Raw 타입 유지: int16, uint8 등 원본 그대로 전송 → 대역폭 절약, PC에서 스케일링")
    add_bullet(doc, "System-Managed: 사용자 코드 없이 자동 전송 (Module ID 0x20)")
    add_bullet(doc, "PhAI V2 호환: 기존 패킷 포맷 변경 없음 (SOF=0xAA, COBS+CRC16)")
    add_bullet(doc, "YAML SSOT: Data Map을 YAML 파일 하나로 관리, Code-Gen으로 C/TS 자동 동기화")

    doc.add_heading("1.3  기존 방식 vs Total Data Packet", level=2)
    add_table(doc,
        ["항목", "기존 (COMBINED 0x10)", "신규 (Total Data 0x20)"],
        [
            ["관리 주체", "User 코드에서 설정", "System 자동 전송"],
            ["데이터 타입", "float32 고정 (10개)", "Raw 타입 유지 (int16/uint8 혼합)"],
            ["패킷 크기", "~54B (COBS)", "~440B (COBS)"],
            ["채널 수", "10개 고정", "87+ 채널 (자동 확장)"],
            ["PC 디코딩", "단순 float 파싱", "Data Map 기반 mixed-type 파싱"],
            ["센서 추가 시", "User 코드 + Studio 양측 수정", "YAML 수정 → 양측 자동 생성"],
            ["사용자 데이터", "같은 패킷에 혼재", "별도 Module ID (0xF0~0xFE)"],
        ], first_col_bold=True
    )

    # ========== Section 2 ==========
    doc.add_heading("2. 설계 결정 요약", level=1)
    add_table(doc,
        ["항목", "결정", "이유"],
        [
            ["전송 모델", "Stream-All (Full Broadcast)", "USB FS 대역폭 충분, XM FW 단순화"],
            ["데이터 타입", "Raw (int16/uint8 유지)", "float 변환 시 2배 크기, 불필요한 연산 제거"],
            ["Data Map 동기화", "YAML → Code-Gen (SSOT)", "양측 자동 동기화, 수동 에러 방지"],
            ["User Custom", "Module ID 0xF0~0xFE (float[])", "Total Data와 공존, 선택적 전송"],
            ["User Meta", "Module ID 0xEF (JSON string)", "PC에서 자동 채널 라벨링"],
            ["CDC 관리 주체", "System (User 코드 불필요)", "사용자 부담 완전 제거"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] 기술사양서 v1.1 Section 3.2: Host Subscribe Model 채택 이유")
    add_ref(doc, "[Ref] 대역폭 계산: USB FS 12Mbps 중 ~425KB/s = 53.1% 점유 (70% 안전 마진 이내)")

    # ========== Section 3 ==========
    doc.add_heading("3. Module ID 체계", level=1)
    add_body(doc,
        "PhAI V2 패킷의 MODULE_ID 필드를 활용하여 데이터 종류를 구분한다. "
        "동일 USB CDC 스트림 위에 COBS 0x00 딜리미터로 복수 패킷이 공존한다."
    )
    add_table(doc,
        ["Module ID", "용도", "주체", "전송 주기"],
        [
            ["0x01~0x07", "Legacy 개별 센서", "—", "Deprecated"],
            ["0x10", "COMBINED (기존)", "User", "Deprecated"],
            ["0x20", "Total Data Packet", "System 자동", "매 1ms"],
            ["0xEF", "User Custom Metadata (JSON)", "System (User 등록)", "연결 시 1회"],
            ["0xF0~0xFE", "User Custom Data (float[])", "User 선택", "User_Loop에서"],
            ["0xFF", "Debug Text", "User 선택", "On-demand"],
        ], first_col_bold=True
    )
    add_code_block(doc,
        "USB CDC 스트림 (1ms 주기):\n"
        "  ├─ [0x20] Total Data ← System 자동 (항상)\n"
        "  ├─ [0xEF] User Meta  ← System (연결 시 1회, User 등록)\n"
        "  └─ [0xF0] User Data  ← User 선택 (User_Loop에서)"
    )

    # ========== Section 4 ==========
    doc.add_heading("4. Total Data Packet 구조", level=1)

    doc.add_heading("4.1  PhAI V2 패킷 래핑", level=2)
    add_code_block(doc,
        "PhAI V2 Packet (기존 포맷 그대로):\n"
        "[SOF:0xAA] [LEN] [SEQ_ID:2B] [MODULE_ID:0x20] [STATUS] [PAYLOAD] [CRC16:2B]\n"
        "                                                         ↓\n"
        "                                              XM_TotalDataPacket_t (~425B)"
    )

    doc.add_heading("4.2  Payload 섹션별 구조", level=2)
    add_table(doc,
        ["#", "섹션", "Offset", "Size", "소스"],
        [
            ["1", "Header", "0", "6B", "System (tick + sensor_mask)"],
            ["2", "CM PDO Raw", "6", "67B", "CM_PdoRx_CmToXm_t memcpy"],
            ["3", "CM SDO", "73", "5B", "CM_RxSdoData_t 개별 복사"],
            ["4", "GRF Left", "78", "18B", "MarvelDex_GetLatest(LEFT)"],
            ["5", "GRF Right", "96", "18B", "MarvelDex_GetLatest(RIGHT)"],
            ["6", "External IMU", "114", "41B", "XsensMTi_GetLatest(0)"],
            ["7", "IMU Hub", "155", "125B", "ImuHub_Drv_GetRxData()"],
            ["8", "External IO", "280", "25B", "DIO bitpack + ADC 12ch"],
            ["9", "Reserved EMG", "305", "40B", "미래 EMG 모듈"],
            ["10", "Reserved HMMG", "345", "40B", "미래 HMMG 모듈"],
            ["11", "Reserved FES", "385", "40B", "미래 FES 모듈"],
            ["", "합계", "", "425B", ""],
        ], first_col_bold=True
    )

    doc.add_heading("4.3  Header 섹션 (6B)", level=2)
    add_table(doc,
        ["필드", "Type", "설명"],
        [
            ["timestamp_ms", "uint32", "System uptime (ms)"],
            ["sensor_mask", "uint16", "연결 센서 비트플래그"],
        ], first_col_bold=True
    )
    add_body(doc, "sensor_mask 비트 정의:")
    add_bullet(doc, "bit 0: CM (H10) 연결")
    add_bullet(doc, "bit 1: GRF Left 연결")
    add_bullet(doc, "bit 2: GRF Right 연결")
    add_bullet(doc, "bit 3: External IMU 연결")
    add_bullet(doc, "bit 4: IMU Hub 연결")
    add_bullet(doc, "bit 5~15: Reserved (향후 EMG/HMMG/FES 등)")

    doc.add_heading("4.4  CM PDO Raw 섹션 (67B)", level=2)
    add_body(doc,
        "CM_PdoRx_CmToXm_t 구조체를 raw memcpy로 전송. PC에서 Data Map 기반 스케일링 적용."
    )
    add_body(doc, "스케일링 공식: physical_value = raw_int16 * SCALE_FACTOR / 32768", bold=True)
    add_table(doc,
        ["필드", "Type", "Scale", "단위", "물리 범위"],
        [
            ["suitAssistModeLoopCnt", "uint32", "1 (no scale)", "count", "0 ~ 4.3B"],
            ["leftHipAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["rightHipAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["leftThighAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["rightThighAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["pelvicAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["pelvicVelY", "int16", "6000", "deg/s", "-3000 ~ +3000"],
            ["leftKneeAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["rightKneeAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["isLeftFootContact", "uint8", "—", "bool", "0/1"],
            ["isRightFootContact", "uint8", "—", "bool", "0/1"],
            ["gaitState", "uint8", "—", "enum", "0~255"],
            ["gaitCycle", "uint8", "—", "%", "0~100"],
            ["forwardVelocity", "int16", "6000", "deg/s", "-3000 ~ +3000"],
            ["leftHipTorque", "int16", "60", "A", "-30 ~ +30"],
            ["rightHipTorque", "int16", "60", "A", "-30 ~ +30"],
            ["leftHipMotorAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["rightHipMotorAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["leftHipImuFrontalRoll", "int16", "720", "deg", "-360 ~ +360"],
            ["leftHipImuSagittalPitch", "int16", "720", "deg", "-360 ~ +360"],
            ["rightHipImuFrontalRoll", "int16", "720", "deg", "-360 ~ +360"],
            ["rightHipImuSagittalPitch", "int16", "720", "deg", "-360 ~ +360"],
            ["leftHipImuGlobalAcc X/Y/Z", "int16 x3", "78.4532", "m/s^2", "-39.24 ~ +39.24"],
            ["leftHipImuGlobalGyr X/Y/Z", "int16 x3", "1000", "deg/s", "-500 ~ +500"],
            ["rightHipImuGlobalAcc X/Y/Z", "int16 x3", "78.4532", "m/s^2", "-39.24 ~ +39.24"],
            ["rightHipImuGlobalGyr X/Y/Z", "int16 x3", "1000", "deg/s", "-500 ~ +500"],
            ["h10FSMcurrentState", "uint8", "—", "enum", "0~255"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] cm_drv.c:122-129 — Scale Factor 정의")
    add_ref(doc, "[Ref] cm_drv.c:2065-2070 — _CM_ScaleInt16ToFloat(): raw * scale / 32768")

    doc.add_heading("4.5  CM SDO 섹션 (5B)", level=2)
    add_body(doc, "이벤트 기반 갱신, 매 tick 최신값 포함.")
    add_table(doc,
        ["필드", "Type", "설명"],
        [
            ["suit_mode", "uint8", "0=STANDBY, 1=ASSIST"],
            ["assist_level", "uint8", "보조 레벨 0~255"],
            ["is_pvec_rh_done", "uint8", "우측 고관절 궤적 완료"],
            ["is_pvec_lh_done", "uint8", "좌측 고관절 궤적 완료"],
            ["is_neutral_pos_set", "uint8", "중립 위치 캘리브레이션 완료"],
        ], first_col_bold=True
    )

    doc.add_heading("4.6  GRF 섹션 (좌/우 각 18B = 36B)", level=2)
    add_table(doc,
        ["필드", "Type", "설명"],
        [
            ["connected", "uint8", "연결 상태 (0/1)"],
            ["sensor_data[14]", "uint8 x14", "FSR 14채널 (0~255 raw)"],
            ["battery", "uint8", "배터리 잔량 (%)"],
            ["status", "uint8", "상태 플래그"],
            ["rolling_idx", "uint8", "시퀀스 번호 (0~199)"],
        ], first_col_bold=True
    )

    doc.add_heading("4.7  External IMU 섹션 (41B)", level=2)
    add_body(doc, "Xsens MTi-630 센서 출력. 원본이 float (센서 내부 캘리브레이션 완료).")
    add_table(doc,
        ["필드", "Type", "단위"],
        [
            ["connected", "uint8", "0/1"],
            ["q_w, q_x, q_y, q_z", "float x4", "normalized quaternion"],
            ["acc_x, acc_y, acc_z", "float x3", "m/s^2"],
            ["gyr_x, gyr_y, gyr_z", "float x3", "deg/s"],
        ], first_col_bold=True
    )

    doc.add_heading("4.8  IMU Hub 섹션 (125B)", level=2)
    add_body(doc,
        "6개 IMU 센서, raw int16 전송. 스케일링 공식이 CM PDO와 다름에 주의."
    )
    add_body(doc, "스케일링 공식: physical = raw / SCALE_FACTOR", bold=True)
    add_table(doc,
        ["필드", "Type", "Scale", "단위"],
        [
            ["timestamp", "uint32", "—", "tick"],
            ["connected_mask", "uint8", "—", "bit0~5 = 센서 0~5"],
            ["imu[0~5].q[4]", "int16 x4", "/10000", "normalized"],
            ["imu[0~5].a[3]", "int16 x3", "/100", "g"],
            ["imu[0~5].g[3]", "int16 x3", "/10", "deg/s"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] imu_hub_drv.h — ImuHub_ImuData_t: int16 q[4]+a[3]+g[3] = 20B/sensor")

    doc.add_heading("4.9  External IO 섹션 (25B)", level=2)
    add_table(doc,
        ["필드", "Type", "설명"],
        [
            ["dio_state", "uint8", "DIO 1~8 bit-packed"],
            ["adc[12]", "uint16 x12", "ADC 12ch (16-bit normalized)"],
        ], first_col_bold=True
    )

    doc.add_heading("4.10  Reserved 섹션 (120B)", level=2)
    add_body(doc, "미래 모듈용 예약 영역. 미사용 시 0 전송.")
    add_table(doc,
        ["섹션", "Size", "용도"],
        [
            ["Reserved EMG", "40B", "향후 EMG 8ch 모듈"],
            ["Reserved HMMG", "40B", "향후 HMMG 모듈"],
            ["Reserved FES", "40B", "향후 FES 모듈"],
        ], first_col_bold=True
    )

    # ========== Section 5 ==========
    doc.add_heading("5. 대역폭 분석", level=1)

    doc.add_heading("5.1  USB CDC Full Speed 제약", level=2)
    add_table(doc,
        ["항목", "값"],
        [
            ["USB FS Raw", "12 Mbps"],
            ["Bulk Packet", "64B max per transaction"],
            ["실효 처리량", "~800~960 KB/s"],
            ["70% 안전 마진", "~560 B/ms"],
        ], first_col_bold=True
    )

    doc.add_heading("5.2  Total Data Packet 대역폭", level=2)
    add_table(doc,
        ["구성", "Payload", "Wire (COBS)", "@ 1kHz", "USB FS 점유"],
        [
            ["Total Data only", "425B", "~443B", "443 KB/s", "55.4%"],
            ["+ User Custom 40B", "465B", "~498B", "498 KB/s", "62.3%"],
            ["70% 마진 이내", "", "", "", "OK"],
        ], first_col_bold=True
    )
    add_body(doc,
        "결론: Total Data(~425B) + User Custom(~40B)을 동시 전송해도 USB FS 대역폭의 62.3%로, "
        "70% 안전 마진 이내에서 안정적 동작이 보장된다.", bold=True
    )
    add_ref(doc, "[Ref] 기술사양서 v1.1 Section 3.2: USB CDC 대역폭 ~464kbps 사용 분석")

    # ========== Section 6 ==========
    doc.add_heading("6. XM FW 데이터 흐름", level=1)
    add_code_block(doc,
        "1ms Tick (StartUserTask)\n"
        "  |\n"
        "  +-- _FetchAllInputs()\n"
        "  |     +-- CM_GetRxData()           -> XM.status.h10 (float, User용)\n"
        "  |     +-- MarvelDex_GetLatest()x2  -> XM.status.grf\n"
        "  |     +-- XsensMTi_GetLatest()     -> XM.status.imu\n"
        "  |     +-- ImuHub_Drv_GetRxData()   -> XM.status.imu_hub\n"
        "  |     +-- XM_IO_Update()\n"
        "  |\n"
        "  +-- XM_TotalData_Snapshot()        <-- NEW: raw 데이터 스냅샷\n"
        "  |     +-- CM_GetRawPdoData()       -> 67B memcpy (int16 raw)\n"
        "  |     +-- CM_GetRxData().sdo       -> 5B 개별 복사\n"
        "  |     +-- MarvelDex raw            -> uint8[14]x2\n"
        "  |     +-- XsensMTi float           -> 40B\n"
        "  |     +-- ImuHub raw int16         -> 120B\n"
        "  |     +-- DIO+ADC                  -> 25B\n"
        "  |\n"
        "  +-- User_Loop()                    <-- 사용자 알고리즘\n"
        "  +-- _FlushAllOutputs()\n"
        "  +-- XM_USB_ProcessPeriodic()\n"
        "        +-- [0xEF] Meta JSON (1회)\n"
        "        +-- [0x20] Total Data (매 tick)\n"
        "        +-- [0xF0] User Custom (선택)"
    )
    add_body(doc,
        "IPO(Input-Process-Output) 모델에서 _FetchAllInputs() 직후 스냅샷을 수행하여, "
        "User_Loop()에서 사용하는 XM.status(float 디코딩)와 Total Data(raw) 간 "
        "시간 일관성을 보장한다."
    )

    # ========== Section 7 ==========
    doc.add_heading("7. Data Map 파이프라인 (YAML → Code-Gen)", level=1)

    doc.add_heading("7.1  SSOT (Single Source of Truth) 개념", level=2)
    add_body(doc,
        "Data Map은 YAML 파일 하나가 SSOT로 동작하며, Python 스크립트가 이를 읽어 "
        "C header와 TypeScript 타입을 자동 생성한다. 양측 코드 수동 동기화가 불필요하다."
    )
    add_code_block(doc,
        "xm_total_data.yaml  -->  generate_data_map.py  -->  xm_total_data_packet.h  (XM FW)\n"
        "                                                -->  xm_total_data_map.ts   (phai-studio)"
    )

    doc.add_heading("7.2  YAML 스키마 포맷", level=2)
    add_code_block(doc,
        "version: \"1.0\"\n"
        "packet_name: XM_TotalDataPacket\n"
        "module_id: 0x20\n"
        "endian: little\n\n"
        "groups:\n"
        "  - name: Header\n"
        "    fields:\n"
        "      - { name: timestamp_ms, type: uint32, unit: ms }\n"
        "      - { name: sensor_mask,  type: uint16, unit: flags }\n"
        "  - name: CM_PDO_Raw\n"
        "    scale_formula: multiply_divide  # raw * scale / 32768\n"
        "    fields:\n"
        "      - { name: leftHipAngle, type: int16, scale: 720, unit: deg }\n"
        "      ...\n\n"
        "types:\n"
        "  imu_hub_sensor_t:\n"
        "    size: 20\n"
        "    fields:\n"
        "      - { name: quat, type: int16, count: 4, scale: 10000, unit: normalized }\n"
        "      - { name: acc,  type: int16, count: 3, scale: 100,   unit: g }\n"
        "      - { name: gyr,  type: int16, count: 3, scale: 10,    unit: deg/s }"
    )

    doc.add_heading("7.3  지원 타입", level=2)
    add_table(doc,
        ["YAML type", "C type", "TS type", "Size"],
        [
            ["uint8", "uint8_t", "'uint8'", "1B"],
            ["int8", "int8_t", "'int8'", "1B"],
            ["uint16", "uint16_t", "'uint16'", "2B"],
            ["int16", "int16_t", "'int16'", "2B"],
            ["uint32", "uint32_t", "'uint32'", "4B"],
            ["int32", "int32_t", "'int32'", "4B"],
            ["float32", "float", "'float32'", "4B"],
        ], first_col_bold=True
    )

    doc.add_heading("7.4  스케일링 공식 (중요)", level=2)
    add_body(doc, "CM PDO와 IMU Hub의 스케일링 공식이 다르므로 YAML에 scale_formula 필드로 구분한다.", bold=True)
    add_table(doc,
        ["공식 ID", "수식", "적용 섹션", "예시"],
        [
            ["multiply_divide", "raw * scale / 32768", "CM PDO Raw",
             "leftHipAngle raw=16384 -> 16384*720/32768 = 360.0 deg"],
            ["divide", "raw / scale", "IMU Hub",
             "quat[0] raw=5000 -> 5000/10000 = 0.5"],
            ["none", "raw (스케일 없음)", "GRF, Header, Reserved",
             "sensor_data raw=128 -> 128"],
        ], first_col_bold=True
    )

    doc.add_heading("7.5  생성 결과물", level=2)
    add_body(doc, "C Header (xm_total_data_packet.h):")
    add_code_block(doc,
        "/* AUTO-GENERATED from xm_total_data.yaml -- DO NOT EDIT MANUALLY */\n"
        "#pragma pack(push, 1)\n"
        "typedef struct {\n"
        "    /* === Header (6B) === */\n"
        "    uint32_t timestamp_ms;           /* offset: 0 */\n"
        "    uint16_t sensor_mask;            /* offset: 4 */\n"
        "    /* === CM_PDO_Raw (67B) === */\n"
        "    uint32_t cm_pdo_loop_cnt;        /* offset: 6 */\n"
        "    int16_t  cm_pdo_left_hip_angle;  /* offset: 10, scale: 720, unit: deg */\n"
        "    // ... (34 fields total)\n"
        "} XM_TotalDataPacket_t;\n"
        "#pragma pack(pop)\n\n"
        "_Static_assert(sizeof(XM_TotalDataPacket_t) == 425,\n"
        "    \"XM_TotalDataPacket_t size mismatch\");\n\n"
        "#define XM_TOTAL_DATA_PAYLOAD_SIZE  sizeof(XM_TotalDataPacket_t)\n"
        "#define XM_TOTAL_DATA_MODULE_ID     0x20\n"
        "#define XM_TOTAL_DATA_NUM_CHANNELS  87"
    )

    add_body(doc, "TypeScript (xm_total_data_map.ts):")
    add_code_block(doc,
        "/* AUTO-GENERATED from xm_total_data.yaml -- DO NOT EDIT MANUALLY */\n\n"
        "export interface ChannelDef {\n"
        "  offset: number;\n"
        "  type: 'uint8'|'int8'|'uint16'|'int16'|'uint32'|'int32'|'float32';\n"
        "  scale: number;\n"
        "  scaleFormula: 'none' | 'divide' | 'multiply_divide';\n"
        "  unit: string;\n"
        "  name: string;\n"
        "  group: string;\n"
        "  count?: number;\n"
        "}\n\n"
        "export const TOTAL_DATA_MAP: ChannelDef[] = [\n"
        "  { offset: 0, type: 'uint32', scale: 1, scaleFormula: 'none',\n"
        "    unit: 'ms', name: 'timestamp_ms', group: 'Header' },\n"
        "  { offset: 10, type: 'int16', scale: 720, scaleFormula: 'multiply_divide',\n"
        "    unit: 'deg', name: 'cm_pdo_left_hip_angle', group: 'CM_PDO' },\n"
        "  // ... (all fields)\n"
        "];\n\n"
        "export const TOTAL_PACKET_SIZE = 425;\n"
        "export const MODULE_ID_TOTAL = 0x20;\n"
        "export const MODULE_ID_USER_META = 0xEF;"
    )
    add_ref(doc, "[Ref] tools/data_map/xm_total_data.yaml — SSOT")
    add_ref(doc, "[Ref] tools/data_map/generate_data_map.py — Code-Gen 스크립트")

    # ========== Section 8 ==========
    doc.add_heading("8. User Custom Data API", level=1)

    doc.add_heading("8.1  개요", level=2)
    add_body(doc,
        "Total Data Packet(0x20)은 System이 자동 전송하므로 사용자 코드가 불필요하다. "
        "사용자가 알고리즘 디버그 데이터(제어 출력, 에러, 커스텀 변수 등)를 추가로 PC에 "
        "전송하고 싶을 때 이 API를 사용한다."
    )

    doc.add_heading("8.2  XM_SetUsbCustomMeta()", level=2)
    add_code_block(doc,
        "/**\n"
        " * @brief User Custom 채널 메타데이터를 등록합니다.\n"
        " *\n"
        " * USB 연결 시 Module ID 0xEF로 자동 전송됩니다.\n"
        " * User_Setup()에서 1회 호출하면 됩니다.\n"
        " *\n"
        " * @param[in] module_id  대상 Module ID (0xF0~0xFE)\n"
        " * @param[in] json_str   채널 정의 JSON string (NULL-terminated)\n"
        " */\n"
        "void XM_SetUsbCustomMeta(uint8_t module_id, const char* json_str);"
    )
    add_body(doc, "JSON 포맷:")
    add_code_block(doc,
        '[\n'
        '  { "name": "Target Angle", "unit": "deg" },\n'
        '  { "name": "Current Angle", "unit": "deg" },\n'
        '  { "name": "Position Error", "unit": "deg" },\n'
        '  { "name": "Torque Output", "unit": "Nm" }\n'
        ']'
    )

    doc.add_heading("8.3  XM_SendUsbDataWithId()", level=2)
    add_code_block(doc,
        "/**\n"
        " * @brief User Custom float[] 데이터를 지정된 Module ID로 전송합니다.\n"
        " *\n"
        " * @param[in] data       float 배열 포인터 (4-byte aligned)\n"
        " * @param[in] len        바이트 수 (sizeof(float) x 채널수)\n"
        " * @param[in] module_id  Module ID (0xF0~0xFE)\n"
        " * @return true: 전송 성공, false: 버퍼 풀 또는 연결 없음\n"
        " */\n"
        "bool XM_SendUsbDataWithId(const void* data, uint32_t len, uint8_t module_id);"
    )

    doc.add_heading("8.4  사용 예시 (Quick Start)", level=2)
    add_code_block(doc,
        '#include "xm_api.h"\n\n'
        "static float s_debug_data[4];\n\n"
        "void User_Setup(void)\n"
        "{\n"
        '    XM_SetUsbCustomMeta(0xF0,\n'
        '        "[{\\"name\\":\\"Target Angle\\",\\"unit\\":\\"deg\\"},"\n'
        '         "{\\"name\\":\\"Current Angle\\",\\"unit\\":\\"deg\\"},"\n'
        '         "{\\"name\\":\\"Position Error\\",\\"unit\\":\\"deg\\"},"\n'
        '         "{\\"name\\":\\"Torque Output\\",\\"unit\\":\\"Nm\\"}]");\n'
        "}\n\n"
        "void User_Loop(void)\n"
        "{\n"
        "    float target = 30.0f;\n"
        "    float current = XM.status.h10.leftHipAngle;\n"
        "    float error = target - current;\n"
        "    float torque = error * 0.5f;  // P-control\n\n"
        "    s_debug_data[0] = target;\n"
        "    s_debug_data[1] = current;\n"
        "    s_debug_data[2] = error;\n"
        "    s_debug_data[3] = torque;\n"
        "    XM_SendUsbDataWithId(s_debug_data, sizeof(s_debug_data), 0xF0);\n\n"
        "    XM_SetAssistTorque(torque, torque);\n"
        "}"
    )

    doc.add_heading("8.5  Wire Format", level=2)
    add_body(doc, "Meta Packet (Module ID 0xEF):")
    add_code_block(doc,
        "PhAI Packet:\n"
        "  [SOF] [LEN] [SEQ_ID] [0xEF] [STATUS]\n"
        "  [PAYLOAD: target_module_id(1B) + json_bytes...]\n"
        "  [CRC16]"
    )
    add_body(doc, "User Data Packet (Module ID 0xF0~0xFE):")
    add_code_block(doc,
        "PhAI Packet:\n"
        "  [SOF] [LEN] [SEQ_ID] [0xF0] [STATUS]\n"
        "  [PAYLOAD: float32[] ...]\n"
        "  [CRC16]"
    )

    doc.add_heading("8.6  제약 조건", level=2)
    add_table(doc,
        ["항목", "제약", "이유"],
        [
            ["Module ID 범위", "0xF0 ~ 0xFE (15개)", "0x20=System, 0xEF=Meta, 0xFF=Debug 예약"],
            ["Payload 타입", "float32 배열 고정", "PC에서 LEN x 4B = float 개수로 자동 디코딩"],
            ["권장 최대 크기", "40 bytes (float x10)", "Total Data(425B)+User(40B)=465B < 560B 안전 마진"],
            ["절대 최대 크기", "120 bytes (float x30)", "이 이상은 USB FS 드롭 위험"],
            ["Meta JSON 크기", "512 bytes 이내", "PhAI LEN 필드 제약 (1020B max)"],
            ["Meta 전송 시점", "연결 시 1회 자동", "재연결 시 자동 재전송"],
        ], first_col_bold=True
    )

    # ========== Section 9 ==========
    doc.add_heading("9. phai-studio 구현 요구사항", level=1)

    add_body(doc,
        "본 섹션은 phai-studio (로컬 에이전트 + 웹 프론트엔드)에서 Total Data Packet을 "
        "수신/디코딩/표시하기 위해 구현해야 할 사항을 상세히 기술한다."
    )

    doc.add_heading("9.1  Module ID 0x20 파서 (Total Data)", level=2)
    add_body(doc, "구현 위치: 로컬 에이전트 (Python) 또는 웹 프론트엔드 (TypeScript)")
    add_body(doc, "필수 구현 사항:", bold=True)
    add_bullet(doc, "MODULE_ID == 0x20 분기 추가 (기존 0x10 COMBINED과 별도)")
    add_bullet(doc, "TOTAL_DATA_MAP (TS 자동 생성 파일) import → 채널 정의 로드")
    add_bullet(doc, "Mixed-type 파싱: ChannelDef의 type 필드에 따라 DataView로 디코딩")
    add_bullet(doc, "스케일링 적용: scaleFormula 필드에 따라 3가지 공식 분기")
    add_bullet(doc, "배열 필드 처리: count > 1인 경우 인덱스 접미사 자동 생성 (예: imu[0].q[0])")

    add_body(doc, "Mixed-Type 파싱 의사코드:")
    add_code_block(doc,
        "function parseTotalData(payload: ArrayBuffer): Map<string, number> {\n"
        "  const view = new DataView(payload);\n"
        "  const result = new Map<string, number>();\n\n"
        "  for (const ch of TOTAL_DATA_MAP) {\n"
        "    const count = ch.count ?? 1;\n"
        "    for (let i = 0; i < count; i++) {\n"
        "      const offset = ch.offset + i * TYPE_SIZE[ch.type];\n"
        "      let raw: number;\n"
        "      switch (ch.type) {\n"
        "        case 'uint8':   raw = view.getUint8(offset); break;\n"
        "        case 'int16':   raw = view.getInt16(offset, true); break;\n"
        "        case 'uint16':  raw = view.getUint16(offset, true); break;\n"
        "        case 'uint32':  raw = view.getUint32(offset, true); break;\n"
        "        case 'float32': raw = view.getFloat32(offset, true); break;\n"
        "        // ...\n"
        "      }\n\n"
        "      // 스케일링 적용\n"
        "      let physical: number;\n"
        "      switch (ch.scaleFormula) {\n"
        "        case 'multiply_divide':\n"
        "          physical = raw * ch.scale / 32768; break;\n"
        "        case 'divide':\n"
        "          physical = raw / ch.scale; break;\n"
        "        case 'none': default:\n"
        "          physical = raw; break;\n"
        "      }\n\n"
        "      const name = count > 1 ? `${ch.name}[${i}]` : ch.name;\n"
        "      result.set(name, physical);\n"
        "    }\n"
        "  }\n"
        "  return result;\n"
        "}"
    )

    doc.add_heading("9.2  Module ID 0xEF 파서 (User Custom Meta)", level=2)
    add_body(doc, "필수 구현 사항:", bold=True)
    add_bullet(doc, "MODULE_ID == 0xEF 수신 시 payload 파싱")
    add_bullet(doc, "payload[0] = target_module_id (0xF0~0xFE), payload[1:] = UTF-8 JSON string")
    add_bullet(doc, "JSON 파싱 → 채널 이름/단위 추출 → 해당 Module ID의 디코딩 메타데이터로 등록")
    add_bullet(doc, "Meta 미수신 시 기본 라벨 사용: \"User_1 Ch0\", \"User_1 Ch1\" ...")

    add_body(doc, "구현 의사코드:")
    add_code_block(doc,
        "function parseUserMeta(payload: Uint8Array): void {\n"
        "  const targetModuleId = payload[0];\n"
        "  const jsonStr = new TextDecoder().decode(payload.slice(1));\n"
        "  const channels = JSON.parse(jsonStr);\n"
        "  // channels: [{ name: 'Target Angle', unit: 'deg' }, ...]\n\n"
        "  userCustomRegistry.set(targetModuleId, channels);\n"
        "}"
    )

    doc.add_heading("9.3  Module ID 0xF0~0xFE 파서 (User Custom Data)", level=2)
    add_body(doc, "필수 구현 사항:", bold=True)
    add_bullet(doc, "MODULE_ID 0xF0~0xFE 범위 수신 시 float32 배열로 디코딩")
    add_bullet(doc, "채널 수 = payload_length / 4 (float32 고정)")
    add_bullet(doc, "Meta 등록된 경우: 채널 이름/단위를 Meta에서 가져와 표시")
    add_bullet(doc, "Meta 미등록 시: 자동 라벨링 (\"User_{id-0xF0+1} Ch{n}\")")

    doc.add_heading("9.4  Channel Picker UI", level=2)
    add_body(doc, "필수 구현 사항:", bold=True)
    add_bullet(doc, "TOTAL_DATA_MAP의 group 필드를 기준으로 그룹별 채널 트리 표시")
    add_bullet(doc, "그룹: Header, CM_PDO, CM_SDO, GRF_Left, GRF_Right, External_IMU, IMU_Hub, External_IO")
    add_bullet(doc, "User Custom 그룹: Module ID 0xEF Meta 수신 시 동적 추가")
    add_bullet(doc, "채널 체크박스: 선택된 채널만 그래프 렌더링에 반영")
    add_bullet(doc, "sensor_mask 연동: bit가 0인 센서 그룹은 UI에서 비활성 표시 (Disconnected)")
    add_bullet(doc, "전체 선택/해제 기능 (그룹 단위)")

    add_body(doc, "UI 레이아웃 예시:")
    add_code_block(doc,
        "Channel Picker:\n"
        "  [v] Header\n"
        "      [v] timestamp_ms (ms)\n"
        "      [v] sensor_mask (flags)\n"
        "  [v] CM_PDO\n"
        "      [v] leftHipAngle (deg)        scale: 720\n"
        "      [v] rightHipAngle (deg)       scale: 720\n"
        "      [ ] leftHipTorque (A)         scale: 60\n"
        "      ...\n"
        "  [ ] GRF_Left  [Disconnected]      <- sensor_mask bit1 = 0\n"
        "      [ ] sensor_data[0~13] (raw)\n"
        "      ...\n"
        "  [v] User Custom (0xF0)\n"
        "      [v] Target Angle (deg)        <- Meta에서 가져온 이름\n"
        "      [v] Current Angle (deg)\n"
        "      [v] Position Error (deg)\n"
        "      [v] Torque Output (Nm)"
    )

    doc.add_heading("9.5  데이터 저장 (Recording)", level=2)
    add_body(doc, "필수 구현 사항:", bold=True)
    add_bullet(doc, "HDF5 또는 CSV 저장 시 Data Map 기반으로 채널명/단위를 메타데이터에 포함")
    add_bullet(doc, "Raw 값 저장 + 스케일 정보 저장 (후처리 시 재스케일링 가능)")
    add_bullet(doc, "또는 스케일링 적용된 물리값으로 저장 (사용자 설정)")
    add_bullet(doc, "User Custom 데이터: 별도 그룹으로 저장 (Module ID별)")
    add_bullet(doc, "sensor_mask 변화 기록: 센서 연결/해제 이벤트 로깅")

    doc.add_heading("9.6  Tx Drop Warning 연동", level=2)
    add_body(doc,
        "PhAI V2 STATUS byte의 bit 0~6 (Tx drop count delta)를 활용하여 "
        "Total Data Packet의 전송 손실을 실시간 모니터링한다."
    )
    add_bullet(doc, "STATUS byte에서 drop_count 추출 (기존 프로토콜 그대로)")
    add_bullet(doc, "drop_count > 0 시 UI에 경고 표시 (기술사양서 v1.1 Section 5 참조)")
    add_bullet(doc, "User Custom 추가로 대역폭 초과 시 drop 증가 → 사용자에게 채널 수 줄이기 안내")

    doc.add_heading("9.7  Data Map 버전 관리", level=2)
    add_body(doc, "필수 구현 사항:", bold=True)
    add_bullet(doc, "xm_total_data_map.ts는 Code-Gen으로 자동 생성 → phai-studio에 import")
    add_bullet(doc, "YAML 버전 필드 (version: \"1.0\")를 TS에도 포함 → 버전 불일치 감지")
    add_bullet(doc, "향후: XM FW가 Data Map 버전을 Handshake 시 전송 → 자동 호환성 검증")
    add_bullet(doc, "TOTAL_PACKET_SIZE와 실제 수신 payload 길이 불일치 시 에러 표시")

    # ========== Section 10 ==========
    doc.add_heading("10. API 변경 요약", level=1)

    doc.add_heading("10.1  신규 API", level=2)
    add_table(doc,
        ["API", "용도"],
        [
            ["XM_SendUsbDataWithId(data, len, module_id)", "User Custom float[] 전송"],
            ["XM_SetUsbCustomMeta(module_id, json_str)", "User Meta 등록 (연결 시 자동 전송)"],
            ["XM_TotalData_Snapshot()", "내부: raw 데이터 스냅샷 수집"],
            ["XM_TotalData_GetLatest()", "내부: 최신 스냅샷 포인터 반환"],
            ["CM_GetRawPdoData()", "내부: CM PDO raw int16 접근"],
        ], first_col_bold=True
    )

    doc.add_heading("10.2  Deprecated API", level=2)
    add_table(doc,
        ["Deprecated API", "대체", "비고"],
        [
            ["XM_SetUsbStreamSource(ptr, size)", "불필요 (System 자동)", "Total Data가 대체"],
            ["XM_SetUsbStreamModuleId(id)", "XM_SendUsbDataWithId()", "Module ID 매 호출 지정"],
            ["XM_SendUsbData(data, len)", "XM_SendUsbDataWithId(data, len, id)", "Module ID 필수"],
        ], first_col_bold=True
    )
    add_body(doc,
        "기존 코드가 Deprecated API를 사용하면 Total Data(0x20)와 함께 이중 전송됩니다. "
        "다음 메이저 버전에서 Deprecated API 제거 예정.", bold=True
    )

    # ========== Section 11 ==========
    doc.add_heading("11. 리스크 및 제약 조건", level=1)
    add_table(doc,
        ["리스크", "영향", "대응"],
        [
            ["pyyaml 미설치", "빌드 실패", "스크립트에 import 체크 + pip install 안내 메시지"],
            ["패킷 >1020B", "LEN 오버플로", "_Static_assert + YAML 크기 계산 (~425B, 여유)"],
            ["CM Mutex 추가 부하", "_FetchAllInputs 지연",
             "기존 CM_GetRxData와 동일 패턴, <5us, 0.5% 미만 충돌"],
            ["Legacy API 사용 유저", "이중 전송",
             "deprecated 주석 + 과도기 호환 유지"],
            ["CM PDO scale 부정확", "PC 디코딩 오류",
             "구현 시 cm_drv.c 실제 코드 확인 후 YAML 기입"],
            ["Reserved 영역 부족", "향후 센서 추가 불가",
             "EMG/HMMG/FES 각 40B = 120B 예약, 필요 시 확장"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] 구현 계획 상세: .claude/plans/splendid-sniffing-finch.md")
    add_ref(doc, "[Ref] 기술사양서 v1.1 Section 3 — Host Subscribe Model 원본")


def build_total_data_packet_doc_en(doc):
    """English version — XM10 Total Data Packet Design Specification."""

    from docx.shared import Pt
    from gen_docx_v2_base import CLR_GRAY

    # === Title Page ===
    add_title_page(doc,
        "XM10 Total Data Packet\nDesign Specification",
        "PhAI Studio Technical Spec v1.1 — Section 3 Host Subscribe Model Implementation",
        "v1.0  ·  2026-03-10  ·  Angel Robotics"
    )

    add_info_cards(doc, [
        ("Reference", "Tech Spec v1.1 Section 3", BG_INFO_BLUE),
        ("Target", "XM10 Rev 1.1 FW", BG_INFO_GRAY),
        ("Packet Size", "~425 Bytes @ 1kHz", BG_INFO_GREEN),
        ("USB FS Usage", "~55.4% (within safety margin)", BG_INFO_GREEN),
    ])
    add_info_cards(doc, [
        ("XM FW", "System auto-transmit (Module ID 0x20)", BG_INFO_BLUE),
        ("User Custom", "Module ID 0xF0~0xFE (optional)", BG_INFO_AMBER),
        ("Data Map", "YAML SSOT → C/TS Code-Gen", BG_INFO_GRAY),
        ("phai-studio", "Phase 2 (deferred)", BG_INFO_GRAY),
    ])

    # === TOC ===
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Overview & Background",
        "2. Design Decisions Summary",
        "3. Module ID Scheme",
        "4. Total Data Packet Structure",
        "5. Bandwidth Analysis",
        "6. XM FW Data Flow",
        "7. Data Map Pipeline (YAML → Code-Gen)",
        "8. User Custom Data API",
        "9. phai-studio Implementation Requirements",
        "10. API Changes Summary",
        "11. Risks & Constraints",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run("   " + item)
        r.font.name = F
        r.font.size = Pt(10)
        r.font.color.rgb = CLR_GRAY
    doc.add_page_break()

    # ========== Section 1 ==========
    doc.add_heading("1. Overview & Background", level=1)

    doc.add_heading("1.1  Relationship to Tech Spec Section 3", level=2)
    add_body(doc,
        "PhAI Studio Technical Specification v1.1, Section 3 \"Data Exchange Model — Host Subscribe Model\" "
        "adopted the Full Broadcast + Host-Side Filtering approach. "
        "This document defines the concrete implementation specification on the XM10 FW side."
    )
    add_body(doc,
        "The original tech spec proposed \"adding fields to the COMBINED packet (Module ID 0x10)\", "
        "which has been evolved into a dedicated Total Data Packet with Module ID 0x20. "
        "This enables the System to automatically transmit all sensor raw data without user code, "
        "while users can select desired channels in phai-studio for monitoring/recording."
    )

    doc.add_heading("1.2  Core Design Principles", level=2)
    add_bullet(doc, "Stream-All, Filter-on-Host: XM sends everything, PC selects")
    add_bullet(doc, "Raw Type Preservation: int16, uint8 etc. transmitted as-is → bandwidth savings, PC-side scaling")
    add_bullet(doc, "System-Managed: Automatic transmission without user code (Module ID 0x20)")
    add_bullet(doc, "PhAI V2 Compatible: No changes to existing packet format (SOF=0xAA, COBS+CRC16)")
    add_bullet(doc, "YAML SSOT: Single YAML file manages Data Map, Code-Gen auto-syncs C/TS")

    doc.add_heading("1.3  Legacy vs Total Data Packet", level=2)
    add_table(doc,
        ["Item", "Legacy (COMBINED 0x10)", "New (Total Data 0x20)"],
        [
            ["Owner", "User code configures", "System auto-transmit"],
            ["Data Type", "float32 only (10 channels)", "Raw types preserved (int16/uint8 mixed)"],
            ["Packet Size", "~54B (COBS)", "~440B (COBS)"],
            ["Channel Count", "10 fixed", "87+ channels (auto-expandable)"],
            ["PC Decoding", "Simple float parsing", "Data Map-based mixed-type parsing"],
            ["Adding Sensors", "User code + Studio both modified", "YAML edit → both auto-generated"],
            ["User Data", "Mixed in same packet", "Separate Module ID (0xF0~0xFE)"],
        ], first_col_bold=True
    )

    # ========== Section 2 ==========
    doc.add_heading("2. Design Decisions Summary", level=1)
    add_table(doc,
        ["Item", "Decision", "Rationale"],
        [
            ["Transmission Model", "Stream-All (Full Broadcast)", "USB FS bandwidth sufficient, simplifies XM FW"],
            ["Data Types", "Raw (int16/uint8 preserved)", "float conversion doubles size, unnecessary computation"],
            ["Data Map Sync", "YAML → Code-Gen (SSOT)", "Auto-sync both sides, prevents manual errors"],
            ["User Custom", "Module ID 0xF0~0xFE (float[])", "Coexists with Total Data, optional transmission"],
            ["User Meta", "Module ID 0xEF (JSON string)", "Automatic channel labeling on PC"],
            ["CDC Management", "System (no user code needed)", "Completely removes user burden"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] Tech Spec v1.1 Section 3.2: Host Subscribe Model adoption rationale")
    add_ref(doc, "[Ref] Bandwidth: USB FS 12Mbps, ~425KB/s = 53.1% utilization (within 70% safety margin)")

    # ========== Section 3 ==========
    doc.add_heading("3. Module ID Scheme", level=1)
    add_body(doc,
        "The MODULE_ID field of PhAI V2 packets is used to distinguish data types. "
        "Multiple packets coexist on the same USB CDC stream, separated by COBS 0x00 delimiters."
    )
    add_table(doc,
        ["Module ID", "Purpose", "Owner", "Frequency"],
        [
            ["0x01~0x07", "Legacy individual sensors", "—", "Deprecated"],
            ["0x10", "COMBINED (legacy)", "User", "Deprecated"],
            ["0x20", "Total Data Packet", "System auto", "Every 1ms"],
            ["0xEF", "User Custom Metadata (JSON)", "System (User registers)", "Once on connect"],
            ["0xF0~0xFE", "User Custom Data (float[])", "User optional", "From User_Loop"],
            ["0xFF", "Debug Text", "User optional", "On-demand"],
        ], first_col_bold=True
    )
    add_code_block(doc,
        "USB CDC Stream (1ms cycle):\n"
        "  +-- [0x20] Total Data  <-- System auto (always)\n"
        "  +-- [0xEF] User Meta   <-- System (once on connect, User registers)\n"
        "  +-- [0xF0] User Data   <-- User optional (from User_Loop)"
    )

    # ========== Section 4 ==========
    doc.add_heading("4. Total Data Packet Structure", level=1)

    doc.add_heading("4.1  PhAI V2 Packet Wrapping", level=2)
    add_code_block(doc,
        "PhAI V2 Packet (existing format unchanged):\n"
        "[SOF:0xAA] [LEN] [SEQ_ID:2B] [MODULE_ID:0x20] [STATUS] [PAYLOAD] [CRC16:2B]\n"
        "                                                         |\n"
        "                                              XM_TotalDataPacket_t (~425B)"
    )

    doc.add_heading("4.2  Payload Section Layout", level=2)
    add_table(doc,
        ["#", "Section", "Offset", "Size", "Source"],
        [
            ["1", "Header", "0", "6B", "System (tick + sensor_mask)"],
            ["2", "CM PDO Raw", "6", "67B", "CM_PdoRx_CmToXm_t memcpy"],
            ["3", "CM SDO", "73", "5B", "CM_RxSdoData_t individual copy"],
            ["4", "GRF Left", "78", "18B", "MarvelDex_GetLatest(LEFT)"],
            ["5", "GRF Right", "96", "18B", "MarvelDex_GetLatest(RIGHT)"],
            ["6", "External IMU", "114", "41B", "XsensMTi_GetLatest(0)"],
            ["7", "IMU Hub", "155", "125B", "ImuHub_Drv_GetRxData()"],
            ["8", "External IO", "280", "25B", "DIO bitpack + ADC 12ch"],
            ["9", "Reserved EMG", "305", "40B", "Future EMG module"],
            ["10", "Reserved HMMG", "345", "40B", "Future HMMG module"],
            ["11", "Reserved FES", "385", "40B", "Future FES module"],
            ["", "Total", "", "425B", ""],
        ], first_col_bold=True
    )

    doc.add_heading("4.3  Header Section (6B)", level=2)
    add_table(doc,
        ["Field", "Type", "Description"],
        [
            ["timestamp_ms", "uint32", "System uptime (ms)"],
            ["sensor_mask", "uint16", "Connected sensor bitflags"],
        ], first_col_bold=True
    )
    add_body(doc, "sensor_mask bit definitions:")
    add_bullet(doc, "bit 0: CM (H10) connected")
    add_bullet(doc, "bit 1: GRF Left connected")
    add_bullet(doc, "bit 2: GRF Right connected")
    add_bullet(doc, "bit 3: External IMU connected")
    add_bullet(doc, "bit 4: IMU Hub connected")
    add_bullet(doc, "bit 5~15: Reserved (future EMG/HMMG/FES etc.)")

    doc.add_heading("4.4  CM PDO Raw Section (67B)", level=2)
    add_body(doc,
        "CM_PdoRx_CmToXm_t struct transmitted via raw memcpy. PC applies Data Map-based scaling."
    )
    add_body(doc, "Scaling formula: physical_value = raw_int16 * SCALE_FACTOR / 32768", bold=True)
    add_table(doc,
        ["Field", "Type", "Scale", "Unit", "Physical Range"],
        [
            ["suitAssistModeLoopCnt", "uint32", "1 (no scale)", "count", "0 ~ 4.3B"],
            ["leftHipAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["rightHipAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["leftThighAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["rightThighAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["pelvicAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["pelvicVelY", "int16", "6000", "deg/s", "-3000 ~ +3000"],
            ["leftKneeAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["rightKneeAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["isLeftFootContact", "uint8", "—", "bool", "0/1"],
            ["isRightFootContact", "uint8", "—", "bool", "0/1"],
            ["gaitState", "uint8", "—", "enum", "0~255"],
            ["gaitCycle", "uint8", "—", "%", "0~100"],
            ["forwardVelocity", "int16", "6000", "deg/s", "-3000 ~ +3000"],
            ["leftHipTorque", "int16", "60", "A", "-30 ~ +30"],
            ["rightHipTorque", "int16", "60", "A", "-30 ~ +30"],
            ["leftHipMotorAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["rightHipMotorAngle", "int16", "720", "deg", "-360 ~ +360"],
            ["leftHipImuFrontalRoll", "int16", "720", "deg", "-360 ~ +360"],
            ["leftHipImuSagittalPitch", "int16", "720", "deg", "-360 ~ +360"],
            ["rightHipImuFrontalRoll", "int16", "720", "deg", "-360 ~ +360"],
            ["rightHipImuSagittalPitch", "int16", "720", "deg", "-360 ~ +360"],
            ["leftHipImuGlobalAcc X/Y/Z", "int16 x3", "78.4532", "m/s^2", "-39.24 ~ +39.24"],
            ["leftHipImuGlobalGyr X/Y/Z", "int16 x3", "1000", "deg/s", "-500 ~ +500"],
            ["rightHipImuGlobalAcc X/Y/Z", "int16 x3", "78.4532", "m/s^2", "-39.24 ~ +39.24"],
            ["rightHipImuGlobalGyr X/Y/Z", "int16 x3", "1000", "deg/s", "-500 ~ +500"],
            ["h10FSMcurrentState", "uint8", "—", "enum", "0~255"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] cm_drv.c:122-129 — Scale Factor definitions")
    add_ref(doc, "[Ref] cm_drv.c:2065-2070 — _CM_ScaleInt16ToFloat(): raw * scale / 32768")

    doc.add_heading("4.5  CM SDO Section (5B)", level=2)
    add_body(doc, "Event-driven updates, latest values included every tick.")
    add_table(doc,
        ["Field", "Type", "Description"],
        [
            ["suit_mode", "uint8", "0=STANDBY, 1=ASSIST"],
            ["assist_level", "uint8", "Assist level 0~255"],
            ["is_pvec_rh_done", "uint8", "Right hip trajectory completed"],
            ["is_pvec_lh_done", "uint8", "Left hip trajectory completed"],
            ["is_neutral_pos_set", "uint8", "Neutral position calibration done"],
        ], first_col_bold=True
    )

    doc.add_heading("4.6  GRF Section (Left/Right each 18B = 36B)", level=2)
    add_table(doc,
        ["Field", "Type", "Description"],
        [
            ["connected", "uint8", "Connection status (0/1)"],
            ["sensor_data[14]", "uint8 x14", "FSR 14ch (0~255 raw)"],
            ["battery", "uint8", "Battery level (%)"],
            ["status", "uint8", "Status flags"],
            ["rolling_idx", "uint8", "Sequence number (0~199)"],
        ], first_col_bold=True
    )

    doc.add_heading("4.7  External IMU Section (41B)", level=2)
    add_body(doc, "Xsens MTi-630 sensor output. Native float (sensor-calibrated internally).")
    add_table(doc,
        ["Field", "Type", "Unit"],
        [
            ["connected", "uint8", "0/1"],
            ["q_w, q_x, q_y, q_z", "float x4", "normalized quaternion"],
            ["acc_x, acc_y, acc_z", "float x3", "m/s^2"],
            ["gyr_x, gyr_y, gyr_z", "float x3", "deg/s"],
        ], first_col_bold=True
    )

    doc.add_heading("4.8  IMU Hub Section (125B)", level=2)
    add_body(doc,
        "6 IMU sensors, raw int16 transmission. Note: scaling formula differs from CM PDO."
    )
    add_body(doc, "Scaling formula: physical = raw / SCALE_FACTOR", bold=True)
    add_table(doc,
        ["Field", "Type", "Scale", "Unit"],
        [
            ["timestamp", "uint32", "—", "tick"],
            ["connected_mask", "uint8", "—", "bit0~5 = sensor 0~5"],
            ["imu[0~5].q[4]", "int16 x4", "/10000", "normalized"],
            ["imu[0~5].a[3]", "int16 x3", "/100", "g"],
            ["imu[0~5].g[3]", "int16 x3", "/10", "deg/s"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] imu_hub_drv.h — ImuHub_ImuData_t: int16 q[4]+a[3]+g[3] = 20B/sensor")

    doc.add_heading("4.9  External IO Section (25B)", level=2)
    add_table(doc,
        ["Field", "Type", "Description"],
        [
            ["dio_state", "uint8", "DIO 1~8 bit-packed"],
            ["adc[12]", "uint16 x12", "ADC 12ch (16-bit normalized)"],
        ], first_col_bold=True
    )

    doc.add_heading("4.10  Reserved Sections (120B)", level=2)
    add_body(doc, "Reserved for future modules. Transmitted as zeroes when unused.")
    add_table(doc,
        ["Section", "Size", "Purpose"],
        [
            ["Reserved EMG", "40B", "Future EMG 8ch module"],
            ["Reserved HMMG", "40B", "Future HMMG module"],
            ["Reserved FES", "40B", "Future FES module"],
        ], first_col_bold=True
    )

    # ========== Section 5 ==========
    doc.add_heading("5. Bandwidth Analysis", level=1)

    doc.add_heading("5.1  USB CDC Full Speed Constraints", level=2)
    add_table(doc,
        ["Item", "Value"],
        [
            ["USB FS Raw", "12 Mbps"],
            ["Bulk Packet", "64B max per transaction"],
            ["Effective Throughput", "~800~960 KB/s"],
            ["70% Safety Margin", "~560 B/ms"],
        ], first_col_bold=True
    )

    doc.add_heading("5.2  Total Data Packet Bandwidth", level=2)
    add_table(doc,
        ["Configuration", "Payload", "Wire (COBS)", "@ 1kHz", "USB FS Usage"],
        [
            ["Total Data only", "425B", "~443B", "443 KB/s", "55.4%"],
            ["+ User Custom 40B", "465B", "~498B", "498 KB/s", "62.3%"],
            ["Within 70% margin", "", "", "", "OK"],
        ], first_col_bold=True
    )
    add_body(doc,
        "Conclusion: Even with Total Data (~425B) + User Custom (~40B) simultaneous transmission, "
        "USB FS bandwidth usage is 62.3%, well within the 70% safety margin.", bold=True
    )
    add_ref(doc, "[Ref] Tech Spec v1.1 Section 3.2: USB CDC bandwidth ~464kbps usage analysis")

    # ========== Section 6 ==========
    doc.add_heading("6. XM FW Data Flow", level=1)
    add_code_block(doc,
        "1ms Tick (StartUserTask)\n"
        "  |\n"
        "  +-- _FetchAllInputs()\n"
        "  |     +-- CM_GetRxData()           -> XM.status.h10 (float, for User)\n"
        "  |     +-- MarvelDex_GetLatest()x2  -> XM.status.grf\n"
        "  |     +-- XsensMTi_GetLatest()     -> XM.status.imu\n"
        "  |     +-- ImuHub_Drv_GetRxData()   -> XM.status.imu_hub\n"
        "  |     +-- XM_IO_Update()\n"
        "  |\n"
        "  +-- XM_TotalData_Snapshot()        <-- NEW: raw data snapshot\n"
        "  |     +-- CM_GetRawPdoData()       -> 67B memcpy (int16 raw)\n"
        "  |     +-- CM_GetRxData().sdo       -> 5B individual copy\n"
        "  |     +-- MarvelDex raw            -> uint8[14]x2\n"
        "  |     +-- XsensMTi float           -> 40B\n"
        "  |     +-- ImuHub raw int16         -> 120B\n"
        "  |     +-- DIO+ADC                  -> 25B\n"
        "  |\n"
        "  +-- User_Loop()                    <-- User algorithm\n"
        "  +-- _FlushAllOutputs()\n"
        "  +-- XM_USB_ProcessPeriodic()\n"
        "        +-- [0xEF] Meta JSON (once)\n"
        "        +-- [0x20] Total Data (every tick)\n"
        "        +-- [0xF0] User Custom (optional)"
    )
    add_body(doc,
        "In the IPO (Input-Process-Output) model, the snapshot is taken immediately after "
        "_FetchAllInputs() to ensure temporal consistency between XM.status (float decoded, "
        "used by User_Loop) and Total Data (raw)."
    )

    # ========== Section 7 ==========
    doc.add_heading("7. Data Map Pipeline (YAML → Code-Gen)", level=1)

    doc.add_heading("7.1  SSOT (Single Source of Truth) Concept", level=2)
    add_body(doc,
        "The Data Map is managed by a single YAML file as the SSOT. A Python script reads it "
        "and auto-generates both a C header and TypeScript types. Manual code synchronization "
        "is completely eliminated."
    )
    add_code_block(doc,
        "xm_total_data.yaml  -->  generate_data_map.py  -->  xm_total_data_packet.h  (XM FW)\n"
        "                                                -->  xm_total_data_map.ts   (phai-studio)"
    )

    doc.add_heading("7.2  YAML Schema Format", level=2)
    add_code_block(doc,
        "version: \"1.0\"\n"
        "packet_name: XM_TotalDataPacket\n"
        "module_id: 0x20\n"
        "endian: little\n\n"
        "groups:\n"
        "  - name: Header\n"
        "    fields:\n"
        "      - { name: timestamp_ms, type: uint32, unit: ms }\n"
        "      - { name: sensor_mask,  type: uint16, unit: flags }\n"
        "  - name: CM_PDO_Raw\n"
        "    scale_formula: multiply_divide  # raw * scale / 32768\n"
        "    fields:\n"
        "      - { name: leftHipAngle, type: int16, scale: 720, unit: deg }\n"
        "      ...\n\n"
        "types:\n"
        "  imu_hub_sensor_t:\n"
        "    size: 20\n"
        "    fields:\n"
        "      - { name: quat, type: int16, count: 4, scale: 10000, unit: normalized }\n"
        "      - { name: acc,  type: int16, count: 3, scale: 100,   unit: g }\n"
        "      - { name: gyr,  type: int16, count: 3, scale: 10,    unit: deg/s }"
    )

    doc.add_heading("7.3  Supported Types", level=2)
    add_table(doc,
        ["YAML type", "C type", "TS type", "Size"],
        [
            ["uint8", "uint8_t", "'uint8'", "1B"],
            ["int8", "int8_t", "'int8'", "1B"],
            ["uint16", "uint16_t", "'uint16'", "2B"],
            ["int16", "int16_t", "'int16'", "2B"],
            ["uint32", "uint32_t", "'uint32'", "4B"],
            ["int32", "int32_t", "'int32'", "4B"],
            ["float32", "float", "'float32'", "4B"],
        ], first_col_bold=True
    )

    doc.add_heading("7.4  Scaling Formulas (Important)", level=2)
    add_body(doc,
        "CM PDO and IMU Hub use different scaling formulas. "
        "The YAML scale_formula field distinguishes them.", bold=True
    )
    add_table(doc,
        ["Formula ID", "Expression", "Applies To", "Example"],
        [
            ["multiply_divide", "raw * scale / 32768", "CM PDO Raw",
             "leftHipAngle raw=16384 -> 16384*720/32768 = 360.0 deg"],
            ["divide", "raw / scale", "IMU Hub",
             "quat[0] raw=5000 -> 5000/10000 = 0.5"],
            ["none", "raw (no scaling)", "GRF, Header, Reserved",
             "sensor_data raw=128 -> 128"],
        ], first_col_bold=True
    )

    doc.add_heading("7.5  Generated Outputs", level=2)
    add_body(doc, "C Header (xm_total_data_packet.h):")
    add_code_block(doc,
        "/* AUTO-GENERATED from xm_total_data.yaml -- DO NOT EDIT MANUALLY */\n"
        "#pragma pack(push, 1)\n"
        "typedef struct {\n"
        "    /* === Header (6B) === */\n"
        "    uint32_t timestamp_ms;           /* offset: 0 */\n"
        "    uint16_t sensor_mask;            /* offset: 4 */\n"
        "    /* === CM_PDO_Raw (67B) === */\n"
        "    uint32_t cm_pdo_loop_cnt;        /* offset: 6 */\n"
        "    int16_t  cm_pdo_left_hip_angle;  /* offset: 10, scale: 720, unit: deg */\n"
        "    // ... (34 fields total)\n"
        "} XM_TotalDataPacket_t;\n"
        "#pragma pack(pop)\n\n"
        "_Static_assert(sizeof(XM_TotalDataPacket_t) == 425,\n"
        "    \"XM_TotalDataPacket_t size mismatch\");\n\n"
        "#define XM_TOTAL_DATA_PAYLOAD_SIZE  sizeof(XM_TotalDataPacket_t)\n"
        "#define XM_TOTAL_DATA_MODULE_ID     0x20\n"
        "#define XM_TOTAL_DATA_NUM_CHANNELS  87"
    )

    add_body(doc, "TypeScript (xm_total_data_map.ts):")
    add_code_block(doc,
        "/* AUTO-GENERATED from xm_total_data.yaml -- DO NOT EDIT MANUALLY */\n\n"
        "export interface ChannelDef {\n"
        "  offset: number;\n"
        "  type: 'uint8'|'int8'|'uint16'|'int16'|'uint32'|'int32'|'float32';\n"
        "  scale: number;\n"
        "  scaleFormula: 'none' | 'divide' | 'multiply_divide';\n"
        "  unit: string;\n"
        "  name: string;\n"
        "  group: string;\n"
        "  count?: number;\n"
        "}\n\n"
        "export const TOTAL_DATA_MAP: ChannelDef[] = [\n"
        "  { offset: 0, type: 'uint32', scale: 1, scaleFormula: 'none',\n"
        "    unit: 'ms', name: 'timestamp_ms', group: 'Header' },\n"
        "  { offset: 10, type: 'int16', scale: 720, scaleFormula: 'multiply_divide',\n"
        "    unit: 'deg', name: 'cm_pdo_left_hip_angle', group: 'CM_PDO' },\n"
        "  // ... (all fields)\n"
        "];\n\n"
        "export const TOTAL_PACKET_SIZE = 425;\n"
        "export const MODULE_ID_TOTAL = 0x20;\n"
        "export const MODULE_ID_USER_META = 0xEF;"
    )
    add_ref(doc, "[Ref] tools/data_map/xm_total_data.yaml — SSOT")
    add_ref(doc, "[Ref] tools/data_map/generate_data_map.py — Code-Gen script")

    # ========== Section 8 ==========
    doc.add_heading("8. User Custom Data API", level=1)

    doc.add_heading("8.1  Overview", level=2)
    add_body(doc,
        "The Total Data Packet (0x20) is automatically transmitted by the System, requiring no user code. "
        "When users want to additionally send algorithm debug data (control outputs, errors, custom variables) "
        "to the PC, they use this API."
    )

    doc.add_heading("8.2  XM_SetUsbCustomMeta()", level=2)
    add_code_block(doc,
        "/**\n"
        " * @brief Registers User Custom channel metadata.\n"
        " *\n"
        " * Automatically transmitted via Module ID 0xEF on USB connection.\n"
        " * Call once in User_Setup().\n"
        " *\n"
        " * @param[in] module_id  Target Module ID (0xF0~0xFE)\n"
        " * @param[in] json_str   Channel definition JSON string (NULL-terminated)\n"
        " */\n"
        "void XM_SetUsbCustomMeta(uint8_t module_id, const char* json_str);"
    )
    add_body(doc, "JSON Format:")
    add_code_block(doc,
        '[\n'
        '  { "name": "Target Angle", "unit": "deg" },\n'
        '  { "name": "Current Angle", "unit": "deg" },\n'
        '  { "name": "Position Error", "unit": "deg" },\n'
        '  { "name": "Torque Output", "unit": "Nm" }\n'
        ']'
    )

    doc.add_heading("8.3  XM_SendUsbDataWithId()", level=2)
    add_code_block(doc,
        "/**\n"
        " * @brief Sends User Custom float[] data with specified Module ID.\n"
        " *\n"
        " * @param[in] data       float array pointer (4-byte aligned)\n"
        " * @param[in] len        Byte count (sizeof(float) x channel_count)\n"
        " * @param[in] module_id  Module ID (0xF0~0xFE)\n"
        " * @return true: success, false: buffer full or not connected\n"
        " */\n"
        "bool XM_SendUsbDataWithId(const void* data, uint32_t len, uint8_t module_id);"
    )

    doc.add_heading("8.4  Usage Example (Quick Start)", level=2)
    add_code_block(doc,
        '#include "xm_api.h"\n\n'
        "static float s_debug_data[4];\n\n"
        "void User_Setup(void)\n"
        "{\n"
        '    XM_SetUsbCustomMeta(0xF0,\n'
        '        "[{\\"name\\":\\"Target Angle\\",\\"unit\\":\\"deg\\"},"\n'
        '         "{\\"name\\":\\"Current Angle\\",\\"unit\\":\\"deg\\"},"\n'
        '         "{\\"name\\":\\"Position Error\\",\\"unit\\":\\"deg\\"},"\n'
        '         "{\\"name\\":\\"Torque Output\\",\\"unit\\":\\"Nm\\"}]");\n'
        "}\n\n"
        "void User_Loop(void)\n"
        "{\n"
        "    float target = 30.0f;\n"
        "    float current = XM.status.h10.leftHipAngle;\n"
        "    float error = target - current;\n"
        "    float torque = error * 0.5f;  // P-control\n\n"
        "    s_debug_data[0] = target;\n"
        "    s_debug_data[1] = current;\n"
        "    s_debug_data[2] = error;\n"
        "    s_debug_data[3] = torque;\n"
        "    XM_SendUsbDataWithId(s_debug_data, sizeof(s_debug_data), 0xF0);\n\n"
        "    XM_SetAssistTorque(torque, torque);\n"
        "}"
    )

    doc.add_heading("8.5  Wire Format", level=2)
    add_body(doc, "Meta Packet (Module ID 0xEF):")
    add_code_block(doc,
        "PhAI Packet:\n"
        "  [SOF] [LEN] [SEQ_ID] [0xEF] [STATUS]\n"
        "  [PAYLOAD: target_module_id(1B) + json_bytes...]\n"
        "  [CRC16]"
    )
    add_body(doc, "User Data Packet (Module ID 0xF0~0xFE):")
    add_code_block(doc,
        "PhAI Packet:\n"
        "  [SOF] [LEN] [SEQ_ID] [0xF0] [STATUS]\n"
        "  [PAYLOAD: float32[] ...]\n"
        "  [CRC16]"
    )

    doc.add_heading("8.6  Constraints", level=2)
    add_table(doc,
        ["Item", "Constraint", "Rationale"],
        [
            ["Module ID Range", "0xF0 ~ 0xFE (15 IDs)", "0x20=System, 0xEF=Meta, 0xFF=Debug reserved"],
            ["Payload Type", "float32 array only", "PC auto-decodes: LEN x 4B = float count"],
            ["Recommended Max", "40 bytes (float x10)", "Total(425B)+User(40B)=465B < 560B safety margin"],
            ["Absolute Max", "120 bytes (float x30)", "Beyond this risks USB FS drops"],
            ["Meta JSON Size", "Within 512 bytes", "PhAI LEN field limit (1020B max)"],
            ["Meta Timing", "Once on connect (auto)", "Auto-retransmit on reconnection"],
        ], first_col_bold=True
    )

    # ========== Section 9 ==========
    doc.add_heading("9. phai-studio Implementation Requirements", level=1)

    add_body(doc,
        "This section details the implementation requirements for phai-studio "
        "(local agent + web frontend) to receive, decode, and display Total Data Packets."
    )

    doc.add_heading("9.1  Module ID 0x20 Parser (Total Data)", level=2)
    add_body(doc, "Implementation location: Local agent (Python) or web frontend (TypeScript)")
    add_body(doc, "Required implementations:", bold=True)
    add_bullet(doc, "Add MODULE_ID == 0x20 branch (separate from legacy 0x10 COMBINED)")
    add_bullet(doc, "Import TOTAL_DATA_MAP (auto-generated TS file) → load channel definitions")
    add_bullet(doc, "Mixed-type parsing: decode via DataView according to ChannelDef type field")
    add_bullet(doc, "Apply scaling: branch on 3 formulas based on scaleFormula field")
    add_bullet(doc, "Array field handling: auto-generate index suffix when count > 1 (e.g., imu[0].q[0])")

    add_body(doc, "Mixed-Type Parsing Pseudocode:")
    add_code_block(doc,
        "function parseTotalData(payload: ArrayBuffer): Map<string, number> {\n"
        "  const view = new DataView(payload);\n"
        "  const result = new Map<string, number>();\n\n"
        "  for (const ch of TOTAL_DATA_MAP) {\n"
        "    const count = ch.count ?? 1;\n"
        "    for (let i = 0; i < count; i++) {\n"
        "      const offset = ch.offset + i * TYPE_SIZE[ch.type];\n"
        "      let raw: number;\n"
        "      switch (ch.type) {\n"
        "        case 'uint8':   raw = view.getUint8(offset); break;\n"
        "        case 'int16':   raw = view.getInt16(offset, true); break;\n"
        "        case 'uint16':  raw = view.getUint16(offset, true); break;\n"
        "        case 'uint32':  raw = view.getUint32(offset, true); break;\n"
        "        case 'float32': raw = view.getFloat32(offset, true); break;\n"
        "        // ...\n"
        "      }\n\n"
        "      // Apply scaling\n"
        "      let physical: number;\n"
        "      switch (ch.scaleFormula) {\n"
        "        case 'multiply_divide':\n"
        "          physical = raw * ch.scale / 32768; break;\n"
        "        case 'divide':\n"
        "          physical = raw / ch.scale; break;\n"
        "        case 'none': default:\n"
        "          physical = raw; break;\n"
        "      }\n\n"
        "      const name = count > 1 ? `${ch.name}[${i}]` : ch.name;\n"
        "      result.set(name, physical);\n"
        "    }\n"
        "  }\n"
        "  return result;\n"
        "}"
    )

    doc.add_heading("9.2  Module ID 0xEF Parser (User Custom Meta)", level=2)
    add_body(doc, "Required implementations:", bold=True)
    add_bullet(doc, "Parse payload on MODULE_ID == 0xEF reception")
    add_bullet(doc, "payload[0] = target_module_id (0xF0~0xFE), payload[1:] = UTF-8 JSON string")
    add_bullet(doc, "Parse JSON → extract channel names/units → register as decoding metadata for that Module ID")
    add_bullet(doc, "Fallback labels when Meta not received: \"User_1 Ch0\", \"User_1 Ch1\" ...")

    add_body(doc, "Implementation pseudocode:")
    add_code_block(doc,
        "function parseUserMeta(payload: Uint8Array): void {\n"
        "  const targetModuleId = payload[0];\n"
        "  const jsonStr = new TextDecoder().decode(payload.slice(1));\n"
        "  const channels = JSON.parse(jsonStr);\n"
        "  // channels: [{ name: 'Target Angle', unit: 'deg' }, ...]\n\n"
        "  userCustomRegistry.set(targetModuleId, channels);\n"
        "}"
    )

    doc.add_heading("9.3  Module ID 0xF0~0xFE Parser (User Custom Data)", level=2)
    add_body(doc, "Required implementations:", bold=True)
    add_bullet(doc, "Decode as float32 array when MODULE_ID in 0xF0~0xFE range")
    add_bullet(doc, "Channel count = payload_length / 4 (float32 fixed)")
    add_bullet(doc, "If Meta registered: display channel names/units from Meta")
    add_bullet(doc, "If Meta not registered: auto-label (\"User_{id-0xF0+1} Ch{n}\")")

    doc.add_heading("9.4  Channel Picker UI", level=2)
    add_body(doc, "Required implementations:", bold=True)
    add_bullet(doc, "Display grouped channel tree based on TOTAL_DATA_MAP group field")
    add_bullet(doc, "Groups: Header, CM_PDO, CM_SDO, GRF_Left, GRF_Right, External_IMU, IMU_Hub, External_IO")
    add_bullet(doc, "User Custom groups: dynamically added upon Module ID 0xEF Meta reception")
    add_bullet(doc, "Channel checkboxes: only selected channels rendered in graphs")
    add_bullet(doc, "sensor_mask integration: sensor groups with bit=0 shown as disabled (Disconnected)")
    add_bullet(doc, "Select/deselect all functionality (per group)")

    add_body(doc, "UI Layout Example:")
    add_code_block(doc,
        "Channel Picker:\n"
        "  [v] Header\n"
        "      [v] timestamp_ms (ms)\n"
        "      [v] sensor_mask (flags)\n"
        "  [v] CM_PDO\n"
        "      [v] leftHipAngle (deg)        scale: 720\n"
        "      [v] rightHipAngle (deg)       scale: 720\n"
        "      [ ] leftHipTorque (A)         scale: 60\n"
        "      ...\n"
        "  [ ] GRF_Left  [Disconnected]      <- sensor_mask bit1 = 0\n"
        "      [ ] sensor_data[0~13] (raw)\n"
        "      ...\n"
        "  [v] User Custom (0xF0)\n"
        "      [v] Target Angle (deg)        <- name from Meta\n"
        "      [v] Current Angle (deg)\n"
        "      [v] Position Error (deg)\n"
        "      [v] Torque Output (Nm)"
    )

    doc.add_heading("9.5  Data Recording", level=2)
    add_body(doc, "Required implementations:", bold=True)
    add_bullet(doc, "Include channel names/units as metadata when saving to HDF5 or CSV (Data Map-based)")
    add_bullet(doc, "Store raw values + scale info (enables re-scaling during post-processing)")
    add_bullet(doc, "Alternatively, store scaled physical values (user configurable)")
    add_bullet(doc, "User Custom data: store in separate groups (per Module ID)")
    add_bullet(doc, "Log sensor_mask changes: record sensor connect/disconnect events")

    doc.add_heading("9.6  Tx Drop Warning Integration", level=2)
    add_body(doc,
        "Utilize PhAI V2 STATUS byte bits 0~6 (Tx drop count delta) to "
        "monitor Total Data Packet transmission losses in real-time."
    )
    add_bullet(doc, "Extract drop_count from STATUS byte (existing protocol unchanged)")
    add_bullet(doc, "Display UI warning when drop_count > 0 (see Tech Spec v1.1 Section 5)")
    add_bullet(doc, "When User Custom additions cause bandwidth overflow → drops increase → advise user to reduce channels")

    doc.add_heading("9.7  Data Map Version Management", level=2)
    add_body(doc, "Required implementations:", bold=True)
    add_bullet(doc, "xm_total_data_map.ts is auto-generated via Code-Gen → imported by phai-studio")
    add_bullet(doc, "Include YAML version field (version: \"1.0\") in TS → detect version mismatch")
    add_bullet(doc, "Future: XM FW transmits Data Map version during Handshake → automatic compatibility check")
    add_bullet(doc, "Display error when TOTAL_PACKET_SIZE mismatches actual received payload length")

    # ========== Section 10 ==========
    doc.add_heading("10. API Changes Summary", level=1)

    doc.add_heading("10.1  New APIs", level=2)
    add_table(doc,
        ["API", "Purpose"],
        [
            ["XM_SendUsbDataWithId(data, len, module_id)", "Send User Custom float[] data"],
            ["XM_SetUsbCustomMeta(module_id, json_str)", "Register User Meta (auto-sent on connect)"],
            ["XM_TotalData_Snapshot()", "Internal: raw data snapshot collection"],
            ["XM_TotalData_GetLatest()", "Internal: return latest snapshot pointer"],
            ["CM_GetRawPdoData()", "Internal: access CM PDO raw int16 data"],
        ], first_col_bold=True
    )

    doc.add_heading("10.2  Deprecated APIs", level=2)
    add_table(doc,
        ["Deprecated API", "Replacement", "Notes"],
        [
            ["XM_SetUsbStreamSource(ptr, size)", "Not needed (System auto)", "Replaced by Total Data"],
            ["XM_SetUsbStreamModuleId(id)", "XM_SendUsbDataWithId()", "Module ID specified per call"],
            ["XM_SendUsbData(data, len)", "XM_SendUsbDataWithId(data, len, id)", "Module ID now required"],
        ], first_col_bold=True
    )
    add_body(doc,
        "If existing code uses Deprecated APIs, data will be dual-transmitted alongside "
        "Total Data (0x20). Deprecated APIs will be removed in the next major version.", bold=True
    )

    # ========== Section 11 ==========
    doc.add_heading("11. Risks & Constraints", level=1)
    add_table(doc,
        ["Risk", "Impact", "Mitigation"],
        [
            ["pyyaml not installed", "Build failure", "Script checks import + displays pip install instructions"],
            ["Packet >1020B", "LEN overflow", "_Static_assert + YAML size calculation (~425B, ample margin)"],
            ["CM Mutex overhead", "_FetchAllInputs delay",
             "Same pattern as CM_GetRxData, <5us, <0.5% collision rate"],
            ["Legacy API users", "Dual transmission",
             "Deprecated annotations + transitional compatibility maintained"],
            ["CM PDO scale inaccuracy", "PC decoding errors",
             "Verify against cm_drv.c actual code before YAML entry"],
            ["Insufficient reserved space", "Cannot add future sensors",
             "EMG/HMMG/FES each 40B = 120B reserved, expandable if needed"],
        ], first_col_bold=True
    )
    add_ref(doc, "[Ref] Detailed implementation plan: .claude/plans/splendid-sniffing-finch.md")
    add_ref(doc, "[Ref] Tech Spec v1.1 Section 3 — Host Subscribe Model original")


def main():
    # Korean version
    doc_kr = setup_document(FONT_KR)
    build_total_data_packet_doc(doc_kr)
    kr_path = os.path.join(OUTPUT_DIR, "XM10_Total_Data_Packet_설계사양서_v1.1.docx")
    doc_kr.save(kr_path)
    print(f"[OK] Korean: {kr_path}")

    # English version
    doc_en = setup_document(FONT_KR)
    build_total_data_packet_doc_en(doc_en)
    en_path = os.path.join(OUTPUT_DIR, "XM10_Total_Data_Packet_Design_Spec_v1.1.docx")
    doc_en.save(en_path)
    print(f"[OK] English: {en_path}")


if __name__ == "__main__":
    main()